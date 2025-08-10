
import os
import requests
import json
import base64
from io import BytesIO
from pathlib import Path
from pdf2image import convert_from_path
from PIL import Image, ImageFilter
import numpy as np
import cv2
from typing import List, Dict
import time
import subprocess
import sys
from collections import defaultdict

class MistralPixtralOCR:
    def __init__(self, api_key: str):
        """
        Initialize Mistral Pixtral OCR client
        """
        self.api_key = api_key
        self.base_url = "https://api.mistral.ai/v1/chat/completions"  # Fixed URL
        self.model = "pixtral-12b-2409"

        # Headers for API requests
        self.headers = {
            "Content-Type": "application/json",
            "Authorization": f"Bearer {self.api_key}"
        }

        # Ensure required packages are installed
        self.ensure_required_packages()

    def ensure_required_packages(self):
        """Ensure all required packages are installed"""
        required_packages = [
            ("docx", "python-docx"),
            ("numpy", "numpy"),
            ("cv2", "opencv-python")
        ]
        
        for module_name, package_name in required_packages:
            try:
                __import__(module_name)
            except ImportError:
                print(f"Installing {package_name}...")
                subprocess.check_call([sys.executable, "-m", "pip", "install", package_name])
                print(f"✅ {package_name} installed successfully!")

    def test_api_connection(self) -> Dict:
        """
        Test API connection with a simple request
        """
        test_payload = {
            "model": self.model,
            "messages": [{"role": "user", "content": "Hello, test connection"}],
            "max_tokens": 10
        }

        try:
            response = requests.post(
                self.base_url,
                headers=self.headers,
                json=test_payload,
                timeout=30
            )

            if response.status_code == 200:
                return {"success": True, "message": "API connection successful"}
            else:
                return {"success": False, "error": f"API Error {response.status_code}: {response.text}"}

        except Exception as e:
            return {"success": False, "error": f"Connection failed: {str(e)}"}

    def is_empty_page(self, image: Image.Image, threshold: float = 0.99) -> bool:
        """
        Check if a page is empty (mostly white/blank)
        """
        # Convert to grayscale
        gray = image.convert('L')
        
        # Convert to numpy array
        np_image = np.array(gray)
        
        # Calculate percentage of white pixels
        white_pixels = np.sum(np_image > 240)  # Threshold for "white"
        total_pixels = np_image.size
        white_ratio = white_pixels / total_pixels
        
        return white_ratio > threshold

    def detect_diagram_regions(self, image: Image.Image) -> List[tuple]:
        """
        Detect regions that are likely diagrams (non-text areas)
        """
        # Convert PIL to OpenCV format
        opencv_image = cv2.cvtColor(np.array(image), cv2.COLOR_RGB2BGR)
        
        # Convert to grayscale
        gray = cv2.cvtColor(opencv_image, cv2.COLOR_BGR2GRAY)
        
        # Apply threshold to get binary image
        _, binary = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY_INV + cv2.THRESH_OTSU)
        
        # Find contours
        contours, _ = cv2.findContours(binary, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
        
        # Filter contours based on area and aspect ratio
        diagram_regions = []
        image_area = gray.shape[0] * gray.shape[1]
        
        for contour in contours:
            x, y, w, h = cv2.boundingRect(contour)
            area = w * h
            aspect_ratio = w / h if h > 0 else 0
            
            # Consider as diagram if:
            # 1. Area is significant (more than 1% of page)
            # 2. Not too narrow or tall (aspect ratio between 0.2 and 5)
            if area > image_area * 0.01 and 0.2 < aspect_ratio < 5:
                diagram_regions.append((x, y, w, h))
        
        return diagram_regions

    def separate_text_and_diagram_regions(self, image: Image.Image) -> tuple:
        """
        Separate text and diagram regions in an image
        Returns: (text_image, diagram_images_list)
        """
        # First, check if it's an empty page
        if self.is_empty_page(image):
            return image, []  # Return original image and no diagrams
        
        # Detect diagram regions
        diagram_regions = self.detect_diagram_regions(image)
        
        if not diagram_regions:
            # No diagrams found, return original image
            return image, []
        
        # Create text-only image by masking diagram regions
        text_image = image.copy()
        diagram_images = []
        
        # Convert to numpy array for easier manipulation
        np_text_image = np.array(text_image)
        
        # Mask diagram regions in text image and extract diagram images
        for i, (x, y, w, h) in enumerate(diagram_regions):
            # Extract diagram region
            diagram_region = image.crop((x, y, x + w, y + h))
            diagram_images.append(diagram_region)
            
            # Mask the diagram region in text image (fill with white)
            np_text_image[y:y+h, x:x+w] = [255, 255, 255]
        
        # Convert back to PIL Image
        text_image = Image.fromarray(np_text_image)
        
        return text_image, diagram_images

    def pdf_to_images(self, pdf_path: str, dpi: int = 300) -> List[Dict]:
        """
        Convert PDF to images with empty page filtering
        Returns: List of dicts with 'image', 'page_number', 'is_empty'
        """
        print(f"Converting PDF to images: {pdf_path}")
        try:
            images = convert_from_path(pdf_path, dpi=dpi)
            processed_pages = []
            
            for i, image in enumerate(images, 1):
                # Check if page is empty
                is_empty = self.is_empty_page(image)
                
                if not is_empty:
                    processed_pages.append({
                        'image': image,
                        'page_number': i,
                        'is_empty': False
                    })
                    print(f"Page {i}: Content detected")
                else:
                    print(f"Page {i}: Empty page removed")
            
            print(f"Successfully converted {len(images)} pages, {len(processed_pages)} with content")
            return processed_pages
        except Exception as e:
            # Enhanced error message for poppler
            print(f"Error converting PDF: {e}. Make sure 'poppler-utils' is installed.")
            return []

    def image_to_base64(self, image: Image.Image) -> str:
        """Convert PIL Image to base64 string"""
        buffer = BytesIO()
        image.save(buffer, format='PNG')
        img_bytes = buffer.getvalue()
        return base64.b64encode(img_bytes).decode('utf-8')

    def transcribe_image(self, image: Image.Image, custom_prompt: str = None, is_diagram: bool = False) -> Dict:
        """Transcribe single image using Mistral Pixtral API"""
        
        # Convert image to base64
        base64_image = self.image_to_base64(image)

        # Different prompts for text and diagram processing
        if is_diagram:
            prompt = custom_prompt or """
            This is a diagram/image. Extract ALL visible text from this image with maximum accuracy.
            Follow these guidelines:
            1. Extract every visible text element including labels, captions, numbers
            2. Preserve original formatting and layout structure
            3. Include any mathematical formulas, symbols, or special characters
            4. Maintain the spatial relationship of text elements
            5. Do NOT describe the image, only extract text
            
            Provide the extracted text maintaining the document structure.
            """
        else:
            prompt = custom_prompt or """
            Transcribe ALL text from this image with maximum accuracy. Follow these guidelines:
            1. Extract every visible text element including headers, body text, footnotes
            2. Preserve original formatting and layout structure
            3. Maintain paragraph breaks and line spacing
            4. Include lists and any formatted content
            5. Transcribe numbers, dates, and special characters accurately
            6. If text is handwritten, read carefully and indicate if unclear
            7. Process multi-column layouts from left to right
            
            Provide the transcribed text maintaining the document structure.
            """

        # Prepare the request payload
        payload = {
            "model": self.model,
            "messages": [
                {
                    "role": "user",
                    "content": [
                        {
                            "type": "text",
                            "text": prompt
                        },
                        {
                            "type": "image_url",
                            "image_url": f"data:image/png;base64,{base64_image}"
                        }
                    ]
                }
            ],
            "max_tokens": 4000,
            "temperature": 0.1
        }

        try:
            response = requests.post(
                self.base_url,
                headers=self.headers,
                json=payload,
                timeout=60
            )

            if response.status_code == 200:
                result = response.json()
                return {
                    "success": True,
                    "text": result['choices'][0]['message']['content'],
                    "usage": result.get('usage', {}),
                    "model": result.get('model', self.model)
                }
            else:
                error_detail = ""
                try:
                    error_data = response.json()
                    error_detail = error_data.get('message', response.text)
                except:
                    error_detail = response.text

                return {
                    "success": False,
                    "error": f"API Error {response.status_code}: {error_detail}"
                }

        except Exception as e:
            return {
                "success": False,
                "error": f"Request failed: {str(e)}"
            }

    def process_page_with_regions(self, page_data: Dict) -> Dict:
        """
        Process a single page by separating text and diagram regions
        """
        image = page_data['image']
        page_number = page_data['page_number']
        
        print(f"Processing page {page_number}...")
        
        # Separate text and diagram regions
        text_image, diagram_images = self.separate_text_and_diagram_regions(image)
        
        # Process text regions
        text_result = self.transcribe_image(text_image, is_diagram=False)
        
        # Process diagram regions
        diagram_results = []
        for i, diagram_image in enumerate(diagram_images):
            print(f"  Processing diagram {i+1} on page {page_number}...")
            diagram_result = self.transcribe_image(diagram_image, is_diagram=True)
            diagram_results.append({
                'index': i+1,
                'result': diagram_result,
                'image': diagram_image
            })
        
        return {
            'page_number': page_number,
            'text_result': text_result,
            'diagram_results': diagram_results,
            'text_image': text_image,
            'diagram_images': diagram_images
        }

    def create_word_document(self, results: List[Dict], output_path: Path, filename: str) -> Path:
        """Create a Word document from extraction results with enhanced formatting"""
        try:
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.oxml.shared import OxmlElement, qn
            from docx.shared import RGBColor

            doc = Document()

            # Set document margins
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
                section.left_margin = Inches(1)
                section.right_margin = Inches(1)

            # Add title with formatting
            title = doc.add_heading(f'OCR Extraction Results', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Add subtitle
            subtitle = doc.add_heading(f'Document: {filename}', level=1)
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Add metadata section
            doc.add_paragraph()  # Empty line
            meta_heading = doc.add_heading('Document Information', level=2)

            # Create metadata table
            meta_table = doc.add_table(rows=5, cols=2)
            meta_table.style = 'Table Grid'

            # Fill metadata
            meta_table.cell(0, 0).text = 'Extraction Method'
            meta_table.cell(0, 1).text = 'Mistral Pixtral API with Diagram Processing'
            meta_table.cell(1, 0).text = 'Total Pages Processed'
            meta_table.cell(1, 1).text = str(len(results))
            meta_table.cell(2, 0).text = 'Successfully Processed'
            meta_table.cell(2, 1).text = str(sum(1 for r in results if r["text_result"]["success"]))
            meta_table.cell(3, 0).text = 'Total Diagrams Found'
            meta_table.cell(3, 1).text = str(sum(len(r["diagram_results"]) for r in results))
            meta_table.cell(4, 0).text = 'Processing Date'
            meta_table.cell(4, 1).text = time.strftime("%Y-%m-%d %H:%M:%S")

            doc.add_paragraph()  # Empty line
            doc.add_paragraph("-" * 80)

            # Add extracted content
            for result in results:
                page_number = result["page_number"]
                
                # Add page header
                page_heading = doc.add_heading(f'Page {page_number}', level=1)
                page_heading.style = 'Heading 1'
                
                # Add text content
                if result["text_result"]["success"] and result["text_result"]["text"]:
                    text_heading = doc.add_heading('Text Content', level=2)
                    text_heading.style = 'Heading 2'
                    
                    # Add character count info
                    info_para = doc.add_paragraph(f'Characters extracted: {len(result["text_result"]["text"])}')
                    info_para.style = 'Intense Quote'

                    # Process and add text content
                    text_content = result["text_result"]["text"].strip()
                    self.add_formatted_content_to_doc(doc, text_content)
                    
                    doc.add_paragraph()  # Empty line

                elif not result["text_result"]["success"]:
                    error_heading = doc.add_heading('Text Processing Error', level=2)
                    error_para = doc.add_paragraph(f'Error: {result["text_result"].get("error", "Unknown error")}')
                    error_para.style = 'Intense Quote'
                    doc.add_paragraph()  # Empty line

                # Add diagram content
                if result["diagram_results"]:
                    diagram_heading = doc.add_heading('Diagrams and Charts', level=2)
                    diagram_heading.style = 'Heading 2'
                    
                    for diagram_result in result["diagram_results"]:
                        diagram_index = diagram_result["index"]
                        diagram_data = diagram_result["result"]
                        
                        diagram_subheading = doc.add_heading(f'Diagram {diagram_index}', level=3)
                        diagram_subheading.style = 'Heading 3'
                        
                        if diagram_data["success"] and diagram_data["text"]:
                            # Add character count info
                            info_para = doc.add_paragraph(f'Characters extracted: {len(diagram_data["text"])}')
                            info_para.style = 'Intense Quote'
                            
                            # Process and add diagram text content
                            diagram_content = diagram_data["text"].strip()
                            self.add_formatted_content_to_doc(doc, diagram_content)
                            
                            doc.add_paragraph()  # Empty line
                        elif not diagram_data["success"]:
                            error_para = doc.add_paragraph(f'Error extracting diagram text: {diagram_data.get("error", "Unknown error")}')
                            error_para.style = 'Intense Quote'
                            doc.add_paragraph()  # Empty line
                
                # Add page separator except for last page
                if page_number < len(results):
                    doc.add_page_break()

            # Save document
            word_file = output_path / f"{filename}_extracted_advanced.docx"
            doc.save(str(word_file))
            print(f"✅ Advanced Word document created successfully: {word_file}")
            return word_file

        except ImportError as e:
            print(f"❌ python-docx not available: {e}")
            print("Creating fallback text file...")
            return self.create_fallback_document(results, output_path, filename)
        except Exception as e:
            print(f"❌ Error creating Word document: {e}")
            print("Creating fallback text file...")
            return self.create_fallback_document(results, output_path, filename)

    def add_formatted_content_to_doc(self, doc, content: str):
        """Add formatted content to Word document, handling paragraphs"""
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        lines = content.split('\n')
        current_paragraph = ""

        for line in lines:
            line = line.strip()
            
            if line:
                if current_paragraph:
                    current_paragraph += " " + line
                else:
                    current_paragraph = line
            else:
                # Empty line indicates paragraph break
                if current_paragraph:
                    para = doc.add_paragraph(current_paragraph)
                    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    current_paragraph = ""

        # Handle remaining content
        if current_paragraph:
            para = doc.add_paragraph(current_paragraph)
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    def create_fallback_document(self, results: List[Dict], output_path: Path, filename: str) -> Path:
        """Create a formatted text file as fallback"""
        fallback_file = output_path / f"{filename}_extracted_fallback_advanced.txt"

        with open(fallback_file, 'w', encoding='utf-8') as f:
            f.write("="*80 + "\n")
            f.write(f"ADVANCED OCR EXTRACTION RESULTS\n")
            f.write(f"Document: {filename}\n")
            f.write("="*80 + "\n\n")

            f.write("DOCUMENT INFORMATION:\n")
            f.write("-" * 40 + "\n")
            f.write(f"Extraction Method: Mistral Pixtral API with Diagram Processing\n")
            f.write(f"Total Pages Processed: {len(results)}\n")
            f.write(f"Successfully Processed: {sum(1 for r in results if r['text_result']['success'])}\n")
            f.write(f"Total Diagrams Found: {sum(len(r['diagram_results']) for r in results)}\n")
            f.write(f"Processing Date: {time.strftime('%Y-%m-%d %H:%M:%S')}\n\n")

            for result in results:
                page_number = result["page_number"]
                f.write("="*80 + "\n")
                f.write(f"PAGE {page_number}\n")
                f.write("="*80 + "\n\n")
                
                # Text content
                if result["text_result"]["success"] and result["text_result"]["text"]:
                    f.write("TEXT CONTENT:\n")
                    f.write("-" * 20 + "\n")
                    f.write(result["text_result"]["text"])
                    f.write("\n\n")
                elif not result["text_result"]["success"]:
                    f.write("TEXT PROCESSING ERROR:\n")
                    f.write("-" * 25 + "\n")
                    f.write(f"Error: {result['text_result'].get('error', 'Unknown error')}\n\n")
                
                # Diagram content
                if result["diagram_results"]:
                    f.write("DIAGRAMS AND CHARTS:\n")
                    f.write("-" * 22 + "\n")
                    for diagram_result in result["diagram_results"]:
                        diagram_index = diagram_result["index"]
                        diagram_data = diagram_result["result"]
                        f.write(f"\nDIAGRAM {diagram_index}:\n")
                        f.write("~" * 15 + "\n")
                        if diagram_data["success"] and diagram_data["text"]:
                            f.write(diagram_data["text"])
                        else:
                            f.write(f"Error: {diagram_data.get('error', 'Unknown error')}")
                        f.write("\n")
                
                f.write("\n")

        print(f"✅ Advanced fallback text document created: {fallback_file}")
        return fallback_file

    def extract_text_from_pdf(self, pdf_path: str, output_dir: str = "ocr_output") -> Dict:
        """Complete PDF text extraction pipeline with advanced features"""
        print(f"Starting advanced PDF text extraction: {pdf_path}")

        # Test API connection first
        connection_test = self.test_api_connection()
        if not connection_test["success"]:
            return {"success": False, "error": f"API connection failed: {connection_test['error']}"}

        # Create output directory
        output_path = Path(output_dir)
        output_path.mkdir(exist_ok=True)

        # Convert PDF to images (with empty page filtering)
        page_data_list = self.pdf_to_images(pdf_path)
        if not page_data_list:
            return {"success": False, "error": "Failed to convert PDF to images. Check poppler-utils installation."}

        # Process each page with region separation
        extraction_results = []
        total_text = ""
        successful_pages = 0
        total_tokens_used = 0
        total_diagrams = 0

        for page_data in page_data_list:
            page_number = page_data['page_number']
            print(f"Processing page {page_number}/{len(page_data_list)}...")
            
            # Process page with region separation
            result = self.process_page_with_regions(page_data)
            extraction_results.append(result)
            
            # Update statistics
            if result["text_result"]["success"]:
                total_text += f"\n\n{'='*50}\nPAGE {page_number} - TEXT\n{'='*50}\n\n{result['text_result']['text']}"
                successful_pages += 1
                
                # Track token usage
                if 'usage' in result["text_result"]:
                    total_tokens_used += result["text_result"]['usage'].get('total_tokens', 0)
            
            total_diagrams += len(result["diagram_results"])
            
            # Save individual page results
            page_dir = output_path / f"page_{page_number:03d}"
            page_dir.mkdir(exist_ok=True)
            
            # Save text result
            if result["text_result"]["success"] and result["text_result"]["text"]:
                text_file = page_dir / "text_content.txt"
                with open(text_file, 'w', encoding='utf-8') as f:
                    f.write(result["text_result"]["text"])
            
            # Save diagram results
            for diagram_result in result["diagram_results"]:
                diagram_index = diagram_result["index"]
                diagram_data = diagram_result["result"]
                
                if diagram_data["success"] and diagram_data["text"]:
                    diagram_text_file = page_dir / f"diagram_{diagram_index}_text.txt"
                    with open(diagram_text_file, 'w', encoding='utf-8') as f:
                        f.write(diagram_data["text"])
                
                # Save diagram image
                diagram_image_file = page_dir / f"diagram_{diagram_index}.png"
                diagram_result["image"].save(diagram_image_file)
            
            # Rate limiting - be respectful to API
            time.sleep(2)

        # Create advanced Word document
        word_file = self.create_word_document(extraction_results, output_path, Path(pdf_path).stem)

        # Also save complete text as backup
        complete_file = output_path / f"{Path(pdf_path).stem}_complete_advanced.txt"
        with open(complete_file, 'w', encoding='utf-8') as f:
            f.write(total_text)

        # Save results summary
        summary = {
            "pdf_file": pdf_path,
            "total_pages_processed": len(page_data_list),
            "successful_pages": successful_pages,
            "failed_pages": len(page_data_list) - successful_pages,
            "total_diagrams_found": total_diagrams,
            "total_characters": len(total_text),
            "total_tokens_used": total_tokens_used,
            "output_files": {
                "word_document": str(word_file),
                "complete_text": str(complete_file),
                "individual_pages": [str(output_path / f"page_{pd['page_number']:03d}") for pd in page_data_list]
            },
            "page_results": extraction_results,
            "processing_time": time.time()
        }

        summary_file = output_path / f"{Path(pdf_path).stem}_summary_advanced.json"
        with open(summary_file, 'w', encoding='utf-8') as f:
            json.dump(summary, f, indent=2, default=str)

        print(f"\n{'='*60}")
        print(f"ADVANCED EXTRACTION COMPLETE!")
        print(f"{'='*60}")
        print(f"✅ Successfully processed: {successful_pages}/{len(page_data_list)} pages")
        print(f"📊 Total diagrams found: {total_diagrams}")
        print(f"📝 Total characters extracted: {len(total_text):,}")
        print(f"🎯 Total tokens used: {total_tokens_used:,}")
        print(f"📁 Output directory: {output_path}")
        print(f"📄 Advanced Word document: {word_file}")
        print(f"📄 Backup text file: {complete_file}")

        return summary


# Usage example
def main():
    # Your API key (keep this secure!)
    API_KEY = os.getenv("MISTRAL_API_KEY", "eyFSYGAUfsrrDmDVLGaKac5IQmFy1gEH")  # Replace with your actual API key

    # Initialize OCR
    ocr = MistralPixtralOCR(API_KEY)

    # Test connection first
    test_result = ocr.test_api_connection()
    if not test_result["success"]:
        print(f"❌ API connection failed: {test_result['error']}")
        return

    print("✅ API connection successful!")

    # Process PDF
    pdf_file = "Lease - Leven(84344979_1) (1).pdf"

    if not Path(pdf_file).exists():
        print(f"❌ PDF file not found: {pdf_file}")
        return

    # Extract text
    result = ocr.extract_text_from_pdf(pdf_file)

    # Check if the result indicates success before accessing keys
    if result and result.get("success", True) and result.get("successful_pages", 0) > 0:
        print(f"\n🎉 SUCCESS! Processed {result['successful_pages']}/{result['total_pages_processed']} pages")
        print(f"📊 Found {result['total_diagrams_found']} diagrams")
        print(f"📄 Advanced Word document available at: {result['output_files']['word_document']}")
    elif result and not result.get("success", True):
         print(f"❌ Extraction failed: {result.get('error', 'Unknown error during extraction')}")
    else:
        print("❌ No pages were successfully processed")


if __name__ == "__main__":
    main()