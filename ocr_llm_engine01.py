import os
import uuid
import json
from datetime import datetime
from io import BytesIO
import ollama
from pdf2image import convert_from_path
from docx import Document
from docx.shared import Inches
from PIL import Image, ImageStat
import numpy as np
from scipy import ndimage
import psycopg2
from psycopg2.extras import RealDictCursor
import base64

# -------------------------------
# 🔧 Configuration
# -------------------------------
DPI = 200
MAX_PAGES = None
MODEL_NAME = "qwen2.5vl"

# Database Configuration
DB_CONFIG = {
    'host': '127.0.0.1',
    'port': 5432,
    'database': 'LLMAPI',
    'user': 'postgres',
    'password': 'shibin'
}

# Optional: set if poppler not in PATH
# POPPLER_PATH = r"C:\path\to\poppler\Library\bin"

# Prompts
TEXT_PROMPT = """
Extract all text from this image of a lease agreement. Include:
- Printed and handwritten content
- Names, addresses, dates, rent amount, lease term
- Signatures, checkboxes, initials
- Clauses, conditions, special terms

Preserve original formatting, line breaks, and structure as much as possible.
If something is unclear, write [Illegible] or [Unclear].
"""

TABLE_PROMPT = """
Analyze this image and determine if it contains a table. If yes:
1. Output ONLY the table in markdown format (using pipes `|` and dashes `-` for headers).
2. Include all text in correct cells.
3. Preserve merged cells by writing [Merged] or leaving appropriate spacing.
4. If a cell is empty, write [Empty].
5. DO NOT include any extra text, explanations, or markdown outside the table.

If no table is present, return exactly: [No Table Detected]
"""

# General OCR prompt for extracting all text from any image
GENERAL_OCR_PROMPT = """
Extract all readable text from this image. This may include:
- Printed text
- Handwritten notes
- Labels, numbers, annotations
- Text in diagrams, maps, or drawings

Preserve original formatting and line breaks where possible.
If text is unclear, write [Illegible].
Return only the extracted text.
"""

# Detection thresholds
EDGE_DENSITY_THRESHOLD = 50


def is_diagram_page(image, edge_threshold=EDGE_DENSITY_THRESHOLD):
    """Detect if a page is a diagram using edge detection."""
    img_gray = image.convert("L")
    np_img = np.array(img_gray).astype(np.float32)
    sobel_x = ndimage.sobel(np_img, axis=0)
    sobel_y = ndimage.sobel(np_img, axis=1)
    edges = np.hypot(sobel_x, sobel_y)
    edge_density = np.mean(edges)
    return edge_density < edge_threshold


def ocr_page(image, prompt):
    """Send image + prompt to Ollama and return extracted text."""
    img_bytes = BytesIO()
    image.convert("RGB").save(img_bytes, format="JPEG")
    img_bytes.seek(0)
    try:
        response = ollama.chat(
            model=MODEL_NAME,
            messages=[
                {
                    "role": "user",
                    "content": prompt,
                    "images": [img_bytes.getvalue()]
                }
            ]
        )
        return response["message"]["content"].strip()
    except Exception as e:
        print(f"❌ OCR Error: {e}")
        return "[OCR Failed]"
    finally:
        img_bytes.close()


def add_markdown_table_to_doc(doc, markdown_table):
    """Converts a markdown-style table into a Word table."""
    lines = [line.strip() for line in markdown_table.strip().split('\n') if line.strip().startswith('|')]
    if len(lines) < 2:
        return False

    # Remove separator line
    header_line = lines[0]
    data_lines = lines[1:]
    if '|---' in data_lines[0]:
        data_lines = data_lines[1:]

    def split_cells(line):
        return [cell.strip() for cell in line.strip('|').split('|')]

    try:
        headers = split_cells(header_line)
        num_cols = len(headers)
        table = doc.add_table(rows=1, cols=num_cols)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        for i, h in enumerate(headers):
            hdr_cells[i].text = h

        for line in data_lines:
            row_cells = split_cells(line)
            if len(row_cells) != num_cols:
                row_cells += ['[Missing]'] * (num_cols - len(row_cells))
                row_cells = row_cells[:num_cols]
            row = table.add_row().cells
            for i, cell_text in enumerate(row_cells):
                row[i].text = cell_text
        return True
    except Exception as e:
        print(f"❌ Table parsing error: {e}")
        return False


def is_table_page(image):
    """Use LLM to detect if the page contains a table."""
    result = ocr_page(image, TABLE_PROMPT)
    return result, "[No Table Detected]" not in result


def get_db_connection():
    """Get database connection"""
    try:
        conn = psycopg2.connect(**DB_CONFIG)
        return conn
    except Exception as e:
        print(f"❌ Database connection error: {e}")
        raise

def init_database():
    """Initialize database tables"""
    conn = get_db_connection()
    cursor = conn.cursor()
    
    try:
        # Create documents table
        cursor.execute("""
            CREATE TABLE IF NOT EXISTS documents (
                id UUID PRIMARY KEY DEFAULT gen_random_uuid(),
                filename VARCHAR(255) NOT NULL,
                original_filename VARCHAR(255) NOT NULL,
                file_size BIGINT,
                file_path TEXT,
                mime_type VARCHAR(100),
                upload_date TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                processing_status VARCHAR(50) DEFAULT 'pending',
                total_pages INTEGER,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            )
        """)
        
        # Create pages table
        cursor.execute("""
            CREATE TABLE IF NOT EXISTS pages (
                id UUID PRIMARY KEY DEFAULT gen_random_uuid(),
                document_id UUID REFERENCES documents(id) ON DELETE CASCADE,
                page_number INTEGER NOT NULL,
                page_type VARCHAR(50), -- 'text', 'table', 'diagram'
                image_data BYTEA,
                image_path TEXT,
                processing_status VARCHAR(50) DEFAULT 'pending',
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            )
        """)
        
        # Create extracted_content table
        cursor.execute("""
            CREATE TABLE IF NOT EXISTS extracted_content (
                id UUID PRIMARY KEY DEFAULT gen_random_uuid(),
                page_id UUID REFERENCES pages(id) ON DELETE CASCADE,
                document_id UUID REFERENCES documents(id) ON DELETE CASCADE,
                content_type VARCHAR(50), -- 'text', 'table', 'metadata'
                raw_text TEXT,
                processed_text TEXT,
                confidence_score FLOAT,
                extraction_method VARCHAR(100),
                metadata JSONB,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            )
        """)
        
        # Create processing_logs table
        cursor.execute("""
            CREATE TABLE IF NOT EXISTS processing_logs (
                id UUID PRIMARY KEY DEFAULT gen_random_uuid(),
                document_id UUID REFERENCES documents(id) ON DELETE CASCADE,
                page_id UUID REFERENCES pages(id) ON DELETE CASCADE,
                step_name VARCHAR(100),
                status VARCHAR(50),
                message TEXT,
                processing_time_ms INTEGER,
                error_details JSONB,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            )
        """)
        
        # Create chat_sessions table
        cursor.execute("""
            CREATE TABLE IF NOT EXISTS chat_sessions (
                id UUID PRIMARY KEY DEFAULT gen_random_uuid(),
                document_id UUID REFERENCES documents(id) ON DELETE CASCADE,
                session_name VARCHAR(255),
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            )
        """)
        
        # Create chat_messages table
        cursor.execute("""
            CREATE TABLE IF NOT EXISTS chat_messages (
                id UUID PRIMARY KEY DEFAULT gen_random_uuid(),
                session_id UUID REFERENCES chat_sessions(id) ON DELETE CASCADE,
                document_id UUID REFERENCES documents(id) ON DELETE CASCADE,
                message_type VARCHAR(20), -- 'user', 'assistant'
                content TEXT NOT NULL,
                metadata JSONB,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            )
        """)
        
        conn.commit()
        print("✅ Database tables initialized successfully")
        
    except Exception as e:
        conn.rollback()
        print(f"❌ Database initialization error: {e}")
        raise
    finally:
        cursor.close()
        conn.close()

def save_document_to_db(filename, original_filename, file_size, file_path, mime_type):
    """Save document metadata to database"""
    conn = get_db_connection()
    cursor = conn.cursor(cursor_factory=RealDictCursor)
    
    try:
        cursor.execute("""
            INSERT INTO documents (filename, original_filename, file_size, file_path, mime_type)
            VALUES (%s, %s, %s, %s, %s)
            RETURNING id, created_at
        """, (filename, original_filename, file_size, file_path, mime_type))
        
        result = cursor.fetchone()
        conn.commit()
        return result['id'], result['created_at']
        
    except Exception as e:
        conn.rollback()
        print(f"❌ Error saving document to database: {e}")
        raise
    finally:
        cursor.close()
        conn.close()

def save_page_to_db(document_id, page_number, page_type, image_data=None, image_path=None):
    """Save page data to database"""
    conn = get_db_connection()
    cursor = conn.cursor(cursor_factory=RealDictCursor)
    
    try:
        cursor.execute("""
            INSERT INTO pages (document_id, page_number, page_type, image_data, image_path)
            VALUES (%s, %s, %s, %s, %s)
            RETURNING id
        """, (document_id, page_number, page_type, image_data, image_path))
        
        result = cursor.fetchone()
        conn.commit()
        return result['id']
        
    except Exception as e:
        conn.rollback()
        print(f"❌ Error saving page to database: {e}")
        raise
    finally:
        cursor.close()
        conn.close()

def save_extracted_content_to_db(page_id, document_id, content_type, raw_text, processed_text=None, confidence_score=None, extraction_method=None, metadata=None):
    """Save extracted content to database"""
    conn = get_db_connection()
    cursor = conn.cursor()
    
    try:
        cursor.execute("""
            INSERT INTO extracted_content 
            (page_id, document_id, content_type, raw_text, processed_text, confidence_score, extraction_method, metadata)
            VALUES (%s, %s, %s, %s, %s, %s, %s, %s)
        """, (page_id, document_id, content_type, raw_text, processed_text, confidence_score, extraction_method, json.dumps(metadata) if metadata else None))
        
        conn.commit()
        
    except Exception as e:
        conn.rollback()
        print(f"❌ Error saving extracted content to database: {e}")
        raise
    finally:
        cursor.close()
        conn.close()

def log_processing_step(document_id, page_id, step_name, status, message=None, processing_time_ms=None, error_details=None):
    """Log processing step to database"""
    conn = get_db_connection()
    cursor = conn.cursor()
    
    try:
        cursor.execute("""
            INSERT INTO processing_logs 
            (document_id, page_id, step_name, status, message, processing_time_ms, error_details)
            VALUES (%s, %s, %s, %s, %s, %s, %s)
        """, (document_id, page_id, step_name, status, message, processing_time_ms, json.dumps(error_details) if error_details else None))
        
        conn.commit()
        
    except Exception as e:
        conn.rollback()
        print(f"❌ Error logging processing step: {e}")
    finally:
        cursor.close()
        conn.close()

def update_document_status(document_id, status, total_pages=None):
    """Update document processing status"""
    conn = get_db_connection()
    cursor = conn.cursor()
    
    try:
        if total_pages is not None:
            cursor.execute("""
                UPDATE documents 
                SET processing_status = %s, total_pages = %s, updated_at = CURRENT_TIMESTAMP
                WHERE id = %s
            """, (status, total_pages, document_id))
        else:
            cursor.execute("""
                UPDATE documents 
                SET processing_status = %s, updated_at = CURRENT_TIMESTAMP
                WHERE id = %s
            """, (status, document_id))
        
        conn.commit()
        
    except Exception as e:
        conn.rollback()
        print(f"❌ Error updating document status: {e}")
        raise
    finally:
        cursor.close()
        conn.close()

def process_pdf_from_frontend(pdf_path, original_filename, document_id=None):
    """Process PDF uploaded from frontend and store all data in database"""
    start_time = datetime.now()
    
    try:
        # Initialize database if not already done
        init_database()
        
        # Save document to database if not provided
        if document_id is None:
            file_size = os.path.getsize(pdf_path)
            document_id, created_at = save_document_to_db(
                filename=os.path.basename(pdf_path),
                original_filename=original_filename,
                file_size=file_size,
                file_path=pdf_path,
                mime_type='application/pdf'
            )
        
        print(f"📄 Starting OCR processing for document ID: {document_id}")
        log_processing_step(document_id, None, 'start_processing', 'started', f"Processing {original_filename}")
        
        # Update status to processing
        update_document_status(document_id, 'processing')
        
        # Convert PDF to images
        try:
            images = convert_from_path(
                pdf_path,
                dpi=DPI,
                first_page=0,
                last_page=MAX_PAGES,
                # poppler_path=POPPLER_PATH
            )
            print(f"✅ Converted {len(images)} pages to images.")
            
            # Update total pages
            update_document_status(document_id, 'processing', len(images))
            log_processing_step(document_id, None, 'pdf_conversion', 'completed', f"Converted {len(images)} pages")
            
        except Exception as e:
            error_msg = f"Failed to convert PDF. Error: {e}"
            log_processing_step(document_id, None, 'pdf_conversion', 'failed', error_msg, error_details={'error': str(e)})
            update_document_status(document_id, 'failed')
            raise RuntimeError(f"❌ {error_msg}")
        
        # Create Word document for output
        doc = Document()
        doc.add_heading("Document - Extracted Content", 0)
        doc.add_paragraph(f"Source: {original_filename}")
        p = doc.add_paragraph("Processed with Qwen2.5vl: Text, Tables, and Diagram-aware OCR.")
        p.runs[0].italic = True
        
        all_extracted_text = []
        
        # Process each page
        for i, image in enumerate(images):
            page_num = i + 1
            total_pages = len(images)
            page_start_time = datetime.now()
            
            print(f"\n📄 Processing Page {page_num}/{total_pages}...")
            
            # Convert image to bytes for storage
            img_bytes = BytesIO()
            image.save(img_bytes, format="JPEG")
            img_bytes.seek(0)
            image_data = img_bytes.getvalue()
            img_bytes.close()
            
            # Determine page type and process
            page_type = 'text'  # default
            extracted_text = ""
            
            try:
                # Step 1: Check for Table (highest priority)
                table_response, is_table = is_table_page(image)
                if is_table:
                    page_type = 'table'
                    extracted_text = table_response
                    print("📊 Detected table – processing...")
                    
                    # Save page to database
                    page_id = save_page_to_db(document_id, page_num, page_type, image_data)
                    
                    # Save extracted content
                    save_extracted_content_to_db(
                        page_id, document_id, 'table', table_response, 
                        extraction_method='qwen2.5vl_table', 
                        metadata={'is_table': True, 'format': 'markdown'}
                    )
                    
                    # Add to document
                    doc.add_heading(f"Page {page_num} (Table)", level=1)
                    success = add_markdown_table_to_doc(doc, table_response)
                    if not success:
                        doc.add_paragraph("[Failed to parse table structure]")
                    
                # Step 2: Check for Diagram
                elif is_diagram_page(image):
                    page_type = 'diagram'
                    extracted_text = ocr_page(image, GENERAL_OCR_PROMPT)
                    print("🎨 Detected diagram/image page – extracting text...")
                    
                    # Save page to database
                    page_id = save_page_to_db(document_id, page_num, page_type, image_data)
                    
                    # Save extracted content
                    save_extracted_content_to_db(
                        page_id, document_id, 'text', extracted_text,
                        extraction_method='qwen2.5vl_general',
                        metadata={'is_diagram': True}
                    )
                    
                    # Add to document
                    doc.add_heading(f"Page {page_num} (Diagram/Image)", level=1)
                    doc.add_paragraph(extracted_text)
                    
                # Step 3: Regular Text Page
                else:
                    page_type = 'text'
                    extracted_text = ocr_page(image, TEXT_PROMPT)
                    print("📝 Detected text page – running OCR...")
                    
                    # Save page to database
                    page_id = save_page_to_db(document_id, page_num, page_type, image_data)
                    
                    # Save extracted content
                    save_extracted_content_to_db(
                        page_id, document_id, 'text', extracted_text,
                        extraction_method='qwen2.5vl_text',
                        metadata={'is_text_page': True}
                    )
                    
                    # Add to document
                    doc.add_heading(f"Page {page_num} (Text)", level=1)
                    doc.add_paragraph(extracted_text)
                
                # Add image to document
                doc.add_heading("🖼️ Original Page Image", level=2)
                img_stream = BytesIO(image_data)
                doc.add_picture(img_stream, width=Inches(6))
                img_stream.close()
                
                # Add page break (except after last)
                if i < total_pages - 1:
                    doc.add_page_break()
                
                # Collect all text
                all_extracted_text.append(f"Page {page_num}: {extracted_text}")
                
                # Log successful page processing
                page_time = (datetime.now() - page_start_time).total_seconds() * 1000
                log_processing_step(document_id, page_id, 'page_processing', 'completed', 
                                  f"Page {page_num} processed as {page_type}", int(page_time))
                
            except Exception as e:
                # Log page processing error
                page_time = (datetime.now() - page_start_time).total_seconds() * 1000
                error_msg = f"Error processing page {page_num}: {e}"
                log_processing_step(document_id, None, 'page_processing', 'failed', 
                                  error_msg, int(page_time), {'error': str(e), 'page_number': page_num})
                print(f"❌ {error_msg}")
                continue
        
        # Save final document
        output_path = f"outputs/{document_id}_processed.docx"
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        doc.save(output_path)
        
        # Save complete document text
        complete_text = "\n\n".join(all_extracted_text)
        save_extracted_content_to_db(
            None, document_id, 'complete_document', complete_text,
            extraction_method='combined_pages',
            metadata={'total_pages': len(images), 'output_path': output_path}
        )
        
        # Update final status
        update_document_status(document_id, 'completed')
        
        # Log completion
        total_time = (datetime.now() - start_time).total_seconds() * 1000
        log_processing_step(document_id, None, 'complete_processing', 'completed', 
                          f"Document processing completed successfully", int(total_time))
        
        print(f"\n🎉 Success! Document processed and saved to database. Document ID: {document_id}")
        
        return {
            'document_id': str(document_id),
            'status': 'completed',
            'total_pages': len(images),
            'extracted_text': complete_text,
            'output_path': output_path,
            'processing_time_ms': int(total_time)
        }
        
    except Exception as e:
        # Log final error
        if 'document_id' in locals():
            update_document_status(document_id, 'failed')
            total_time = (datetime.now() - start_time).total_seconds() * 1000
            log_processing_step(document_id, None, 'complete_processing', 'failed', 
                              f"Document processing failed: {e}", int(total_time), {'error': str(e)})
        
        print(f"❌ Processing failed: {e}")
        raise

def get_document_from_db(document_id):
    """Retrieve document and its content from database"""
    conn = get_db_connection()
    cursor = conn.cursor(cursor_factory=RealDictCursor)
    
    try:
        # Get document info
        cursor.execute("""
            SELECT d.*, 
                   COUNT(p.id) as page_count,
                   STRING_AGG(ec.raw_text, '\n\n' ORDER BY p.page_number) as complete_text
            FROM documents d
            LEFT JOIN pages p ON d.id = p.document_id
            LEFT JOIN extracted_content ec ON p.id = ec.page_id AND ec.content_type = 'text'
            WHERE d.id = %s
            GROUP BY d.id
        """, (document_id,))
        
        document = cursor.fetchone()
        
        if not document:
            return None
        
        # Get pages with content
        cursor.execute("""
            SELECT p.*, ec.raw_text, ec.content_type, ec.metadata
            FROM pages p
            LEFT JOIN extracted_content ec ON p.id = ec.page_id
            WHERE p.document_id = %s
            ORDER BY p.page_number
        """, (document_id,))
        
        pages = cursor.fetchall()
        
        return {
            'document': dict(document),
            'pages': [dict(page) for page in pages]
        }
        
    except Exception as e:
        print(f"❌ Error retrieving document from database: {e}")
        raise
    finally:
        cursor.close()
        conn.close()

def main():
    """Main function for testing - can be removed in production"""
    # This is for testing purposes only
    test_pdf = "test_document.pdf"
    if os.path.exists(test_pdf):
        result = process_pdf_from_frontend(test_pdf, "test_document.pdf")
        print(f"Processing result: {result}")
    else:
        print("No test PDF found. Use process_pdf_from_frontend() function from your FastAPI backend.")

if __name__ == "__main__":
    main()