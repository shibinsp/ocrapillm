import os
import requests
import json
import base64
from io import BytesIO
from pathlib import Path
from pdf2image import convert_from_path
from PIL import Image
from typing import List, Dict
import time
import subprocess
import sys


class MistralPixtralOCR:
    def __init__(self, api_key: str):
        """
        Initialize Mistral Pixtral OCR client
        """
        self.api_key = api_key
        self.base_url = "https://api.mistral.ai/v1/chat/completions"
        self.model = "pixtral-12b-2409"

        # Headers for API requests
        self.headers = {
            "Content-Type": "application/json",
            "Authorization": f"Bearer {self.api_key}"
        }

        # Ensure python-docx is installed
        self.ensure_docx_installed()

    def ensure_docx_installed(self):
        """Ensure python-docx is installed"""
        try:
            import docx
        except ImportError:
            print("Installing python-docx...")
            subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
            print("✅ python-docx installed successfully!")

    def test_api_connection(self) -> Dict:
        """
        Test API connection with a simple request
        """
        test_payload = {
            "model": self.model,
            "messages": [{"role": "user", "content": "Hello, test connection"}],
            "max_tokens": 10
        }

        try:
            response = requests.post(
                self.base_url,
                headers=self.headers,
                json=test_payload,
                timeout=30
            )

            if response.status_code == 200:
                return {"success": True, "message": "API connection successful"}
            else:
                return {"success": False, "error": f"API Error {response.status_code}: {response.text}"}

        except Exception as e:
            return {"success": False, "error": f"Connection failed: {str(e)}"}

    def pdf_to_images(self, pdf_path: str, dpi: int = 300) -> List[Image.Image]:
        """Convert PDF to images"""
        print(f"Converting PDF to images: {pdf_path}")
        try:
            images = convert_from_path(pdf_path, dpi=dpi)
            print(f"Successfully converted {len(images)} pages")
            return images
        except Exception as e:
            # Enhanced error message for poppler
            print(f"Error converting PDF: {e}. Make sure 'poppler-utils' is installed.")
            return []

    def image_to_base64(self, image: Image.Image) -> str:
        """Convert PIL Image to base64 string"""
        buffer = BytesIO()
        image.save(buffer, format='PNG')
        img_bytes = buffer.getvalue()
        return base64.b64encode(img_bytes).decode('utf-8')

    def transcribe_image(self, image: Image.Image, custom_prompt: str = None) -> Dict:
        """Transcribe single image using Mistral Pixtral API"""
        base64_image = self.image_to_base64(image)
        prompt = custom_prompt or """
        Transcribe ALL text from this image with maximum accuracy. Follow these guidelines:
        1. Extract every visible text element including headers, body text, footnotes
        2. Preserve original formatting and layout structure
        3. Maintain paragraph breaks and line spacing
        4. Include tables, lists, and any formatted content
        5. Transcribe numbers, dates, and special characters accurately
        6. If text is handwritten, read carefully and indicate if unclear
        7. Process multi-column layouts from left to right

        Provide the transcribed text maintaining the document structure.
        """
        payload = {
            "model": self.model,
            "messages": [
                {
                    "role": "user",
                    "content": [
                        {"type": "text", "text": prompt},
                        {"type": "image_url", "image_url": f"data:image/png;base64,{base64_image}"}
                    ]
                }
            ],
            "max_tokens": 4000,
            "temperature": 0.1
        }
        try:
            response = requests.post(
                self.base_url,
                headers=self.headers,
                json=payload,
                timeout=60
            )
            if response.status_code == 200:
                result = response.json()
                return {
                    "success": True,
                    "text": result['choices'][0]['message']['content'],
                    "usage": result.get('usage', {}),
                    "model": result.get('model', self.model)
                }
            else:
                error_detail = ""
                try:
                    error_data = response.json()
                    error_detail = error_data.get('message', response.text)
                except:
                    error_detail = response.text
                return {"success": False, "error": f"API Error {response.status_code}: {error_detail}"}
        except Exception as e:
            return {"success": False, "error": f"Request failed: {str(e)}"}

    def create_word_document(self, results: List[Dict], output_path: Path, filename: str, image_paths: List[Path]) -> Path:
        """Create a Word document from extraction results with original images"""
        try:
            from docx import Document
            from docx.shared import Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            doc = Document()
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
                section.left_margin = Inches(1)
                section.right_margin = Inches(1)

            title = doc.add_heading('OCR Extraction Results', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            subtitle = doc.add_heading(f'Document: {filename}', level=1)
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            doc.add_paragraph()
            doc.add_heading('Document Information', level=2)
            meta_table = doc.add_table(rows=4, cols=2)
            meta_table.style = 'Table Grid'
            meta_table.cell(0, 0).text = 'Extraction Method'
            meta_table.cell(0, 1).text = 'Mistral Pixtral API'
            meta_table.cell(1, 0).text = 'Total Pages'
            meta_table.cell(1, 1).text = str(len(results))
            meta_table.cell(2, 0).text = 'Successfully Processed'
            meta_table.cell(2, 1).text = str(sum(1 for r in results if r["status"] == "success"))
            meta_table.cell(3, 0).text = 'Processing Date'
            meta_table.cell(3, 1).text = time.strftime("%Y-%m-%d %H:%M:%S")
            doc.add_paragraph()

            for i, result in enumerate(results):
                doc.add_heading(f'Page {result["page"]}', level=2)
                
                # Embed the original image for the page
                if i < len(image_paths) and image_paths[i] and image_paths[i].exists():
                    try:
                        doc.add_picture(str(image_paths[i]), width=Inches(6.0))
                    except Exception as e:
                        doc.add_paragraph(f"[Could not embed image: {e}]")

                if result["status"] == "success" and result["text"]:
                    doc.add_paragraph(f'Characters extracted: {len(result["text"])}').style = 'Intense Quote'
                    para = doc.add_paragraph(result["text"].strip())
                    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                elif result["status"] == "error":
                    doc.add_paragraph(f'Error: {result.get("error", "Unknown error")}').style = 'Intense Quote'

                if result["page"] < len(results):
                    doc.add_page_break()

            word_file = output_path / f"{filename}_extracted.docx"
            doc.save(str(word_file))
            print(f"✅ Word document created successfully: {word_file}")
            return word_file

        except ImportError as e:
            print(f"❌ python-docx not available, creating fallback text file: {e}")
            return self.create_fallback_document(results, output_path, filename)
        except Exception as e:
            print(f"❌ Error creating Word document, creating fallback text file: {e}")
            return self.create_fallback_document(results, output_path, filename)

    def create_fallback_document(self, results: List[Dict], output_path: Path, filename: str) -> Path:
        """Create a formatted text file as fallback"""
        fallback_file = output_path / f"{filename}_extracted_fallback.txt"
        with open(fallback_file, 'w', encoding='utf-8') as f:
            f.write(f"{'='*80}\nOCR EXTRACTION RESULTS\nDocument: {filename}\n{'='*80}\n\n")
            for result in results:
                f.write(f"{'='*80}\nPAGE {result['page']}\n{'='*80}\n\n")
                if result["status"] == "success":
                    f.write(result["text"] + "\n\n")
                else:
                    f.write(f"ERROR: {result.get('error', 'Unknown error')}\n\n")
        print(f"✅ Fallback text document created: {fallback_file}")
        return fallback_file

    def extract_text_from_pdf(self, pdf_path: str, output_dir: str = "ocr_output") -> Dict:
        """Complete PDF text extraction pipeline"""
        print(f"Starting PDF text extraction: {pdf_path}")
        connection_test = self.test_api_connection()
        if not connection_test["success"]:
            return {"success": False, "error": f"API connection failed: {connection_test['error']}"}

        output_path = Path(output_dir)
        output_path.mkdir(exist_ok=True)
        pdf_stem = Path(pdf_path).stem

        images = self.pdf_to_images(pdf_path)
        if not images:
            return {"success": False, "error": "Failed to convert PDF to images. Check poppler-utils installation."}

        # Save original page images
        images_output_path = output_path / f"{pdf_stem}_images"
        images_output_path.mkdir(exist_ok=True)
        saved_image_paths = []
        for i, image in enumerate(images, 1):
            image_file_path = images_output_path / f"page_{i:03d}.png"
            try:
                image.save(image_file_path, "PNG")
                saved_image_paths.append(image_file_path)
            except Exception as e:
                print(f"❌ Could not save image for page {i}: {e}")
                saved_image_paths.append(None)

        extraction_results = []
        total_text = ""
        successful_pages = 0
        total_tokens_used = 0

        for i, image in enumerate(images, 1):
            print(f"Processing page {i}/{len(images)}...")
            result = self.transcribe_image(image)
            if result["success"]:
                page_text = result["text"]
                total_text += f"\n\n{'='*50}\nPAGE {i}\n{'='*50}\n\n{page_text}"
                successful_pages += 1
                total_tokens_used += result.get('usage', {}).get('total_tokens', 0)
                extraction_results.append({
                    "page": i, "text": page_text, "status": "success", "usage": result.get("usage", {})
                })
                print(f"✅ Page {i} processed successfully ({len(page_text)} characters)")
            else:
                print(f"❌ Error on page {i}: {result['error']}")
                extraction_results.append({
                    "page": i, "text": "", "status": "error", "error": result["error"]
                })
            time.sleep(2)

        word_file = self.create_word_document(extraction_results, output_path, pdf_stem, saved_image_paths)
        
        complete_file = output_path / f"{pdf_stem}_complete.txt"
        with open(complete_file, 'w', encoding='utf-8') as f:
            f.write(total_text)

        summary = {
            "pdf_file": pdf_path,
            "total_pages": len(images),
            "successful_pages": successful_pages,
            "total_tokens_used": total_tokens_used,
            "output_files": {
                "word_document": str(word_file),
                "complete_text": str(complete_file),
                "original_images_dir": str(images_output_path)
            },
            "page_results": extraction_results,
        }
        summary_file = output_path / f"{pdf_stem}_summary.json"
        with open(summary_file, 'w', encoding='utf-8') as f:
            json.dump(summary, f, indent=2)

        print(f"\n{'='*60}\nEXTRACTION COMPLETE!\n{'='*60}")
        print(f"✅ Successfully processed: {successful_pages}/{len(images)} pages")
        print(f"📄 Word document: {word_file}")
        print(f"🖼️ Original images saved in: {images_output_path}")
        print(f"📄 Backup text file: {complete_file}")
        
        return summary

# Usage example
def main():
    # Your API key (keep this secure!)
    API_KEY = "eyFSYGAUfsrrDmDVLGaKac5IQmFy1gEH"  # Your API key is now set here.

    ocr = MistralPixtralOCR(API_KEY)
    
    test_result = ocr.test_api_connection()
    if not test_result["success"]:
        print(f"❌ API connection failed: {test_result['error']}")
        return
    print("✅ API connection successful!")

    pdf_file = "train_model_pdf.pdf"
    if not Path(pdf_file).exists():
        print(f"❌ PDF file not found: {pdf_file}. Please place it in the same directory.")
        return

    result = ocr.extract_text_from_pdf(pdf_file)
    if result and result.get("successful_pages", 0) > 0:
        print(f"\n🎉 SUCCESS! Processed {result['successful_pages']}/{result['total_pages']} pages.")
    elif result and not result.get("success", True):
        print(f"❌ Extraction failed: {result.get('error', 'Unknown error')}")
    else:
        print("❌ No pages were successfully processed.")

if __name__ == "__main__":
    main()
