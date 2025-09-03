from fastapi import FastAPI, File, UploadFile, HTTPException, Depends, Request, Response
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse, FileResponse
from pydantic import BaseModel
import os
import sys
import uuid
import json
from datetime import datetime
from typing import Optional, List, Optional
import shutil
from pathlib import Path
import asyncio
from concurrent.futures import ThreadPoolExecutor
import logging
import time
import traceback

# Sentry configuration (optional)
try:
    import sentry_sdk
    from sentry_sdk.integrations.fastapi import FastApiIntegration
    from sentry_sdk.integrations.starlette import StarletteIntegration
    SENTRY_AVAILABLE = True
    print("[OK] Sentry SDK imported successfully")
except ImportError as e:
    SENTRY_AVAILABLE = False
    sentry_sdk = None
    FastApiIntegration = None
    StarletteIntegration = None
    print(f"[WARN] Sentry SDK not available - error monitoring disabled: {e}")

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('app.log'),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

# Environment validation
def validate_environment():
    """Validate required environment variables and dependencies"""
    warnings = []
    errors = []
    
    # Check required environment variables
    if not os.getenv("SENTRY_DSN"):
        warnings.append("SENTRY_DSN not set - error monitoring disabled")
    
    # Check database connection
    try:
        import psycopg2
        db_config = {
            # 'host': '127.0.0.1',
            'host': 'localhost',
            'port': 5432,
            'database': 'LLMAPI',
            'user': 'postgres',
            # 'password': 'shibin'
            'password': 'sai'
        }
        conn = psycopg2.connect(**db_config)
        conn.close()
        logger.info("[OK] Database connection validated")
    except Exception as e:
        warnings.append(f"Database connection failed: {e}")
        logger.warning(f"[WARN] Database not available: {e}")
    
    # Check external dependencies
    try:
        import pdf2image
        logger.info("[OK] PDF processing libraries available")
    except ImportError:
        errors.append("pdf2image not installed - PDF processing will fail")
    
    try:
        import cv2
        logger.info("[OK] OpenCV available for image processing")
    except ImportError:
        warnings.append("OpenCV not available - some image processing features may be limited")
    
    # Check API connectivity (non-blocking)
    try:
        import requests
        response = requests.get("https://api.mistral.ai", timeout=5)
        logger.info("[OK] Mistral API endpoint reachable")
    except Exception as e:
        warnings.append(f"Mistral API may not be reachable: {e}")
    
    # Log results
    if errors:
        for error in errors:
            logger.error(f"[ERROR] {error}")
        raise Exception(f"Critical startup errors: {', '.join(errors)}")
    
    if warnings:
        for warning in warnings:
            logger.warning(f"[WARN] {warning}")
    
    logger.info("[OK] Environment validation completed")
    return warnings

# Validate environment on startup
try:
    startup_warnings = validate_environment()
except Exception as e:
    logger.error(f"Startup validation failed: {e}")
    # Continue startup but with limited functionality
    startup_warnings = [str(e)]

# Initialize Sentry if available
if SENTRY_AVAILABLE and sentry_sdk and FastApiIntegration and StarletteIntegration:
    try:
        sentry_sdk.init(
            dsn=os.getenv("SENTRY_DSN", "https://321e46d8d2c2350017f0bc48d670a3c1@o4509915357184000.ingest.de.sentry.io/4509925519720528"),  # Sentry DSN for backend monitoring
            integrations=[
                FastApiIntegration(),
                StarletteIntegration(transaction_style="endpoint"),
            ],
            traces_sample_rate=1.0,  # Capture 100% of transactions for performance monitoring
            profiles_sample_rate=1.0,  # Capture 100% of profiles
            environment=os.getenv("ENVIRONMENT", "development"),
        )
        print("[OK] Sentry monitoring initialized")
    except Exception as e:
        print(f"[WARN] Failed to initialize Sentry: {e}")
else:
    print("[INFO] Sentry monitoring disabled - SDK not available")

# Import OCR engine class with error handling
try:
    from ocr_engine_clean import DatabaseOCR
    logger.info("[OK] OCR engine imported successfully")
    # Initialize OCR processor and database tables on startup
    ocr_processor = DatabaseOCR(api_key="eyFSYGAUfsrrDmDVLGaKac5IQmFy1gEH", init_db=False)
    logger.info("[OK] OCR engine initialized successfully")
except Exception as e:
    logger.error(f"Warning: Could not import ocr_engine_clean: {e}. OCR functionality will be limited.")
    ocr_processor = None
    DatabaseOCR = None

# Import new workflow modules with error handling
try:
    from arc_diagram_separation import ArcDiagramSeparator
    from text_extraction_from_diagram import DiagramTextExtractor
    logger.info("[OK] Workflow modules imported successfully")
except Exception as e:
    logger.error(f"Warning: Could not import workflow modules: {e}. New workflow will be limited.")
    ArcDiagramSeparator = None
    DiagramTextExtractor = None

# Utility functions for error handling
def handle_database_error(operation: str, error: Exception):
    """Centralized database error handling"""
    logger.error(f"Database operation '{operation}' failed: {str(error)}")
    logger.error(f"Traceback: {traceback.format_exc()}")
    return {
        "error": "database_error",
        "message": f"Database operation failed: {operation}",
        "details": str(error) if os.getenv("DEBUG") else "Internal server error"
    }

def handle_api_error(service: str, error: Exception):
    """Centralized external API error handling"""
    logger.error(f"External API '{service}' failed: {str(error)}")
    return {
        "error": "api_error",
        "message": f"External service unavailable: {service}",
        "details": str(error) if os.getenv("DEBUG") else "Service temporarily unavailable"
    }

def retry_with_backoff(func, max_retries: int = 3, base_delay: float = 1.0):
    """Retry function with exponential backoff"""
    for attempt in range(max_retries):
        try:
            return func()
        except Exception as e:
            if attempt == max_retries - 1:
                raise e
            delay = base_delay * (2 ** attempt)
            logger.warning(f"Attempt {attempt + 1} failed, retrying in {delay}s: {str(e)}")
            time.sleep(delay)

app = FastAPI(
    title="OCR AI Assistant API",
    description="Backend API for OCR document processing and AI chatbot",
    version="1.0.0"
)

# Error handling middleware
@app.middleware("http")
async def error_handling_middleware(request: Request, call_next):
    start_time = time.time()
    try:
        response = await call_next(request)
        process_time = time.time() - start_time
        logger.info(f"{request.method} {request.url} - {response.status_code} - {process_time:.3f}s")
        return response
    except Exception as e:
        process_time = time.time() - start_time
        logger.error(f"{request.method} {request.url} - ERROR - {process_time:.3f}s: {str(e)}")
        logger.error(f"Traceback: {traceback.format_exc()}")
        return JSONResponse(
            status_code=500,
            content={
                "error": "internal_server_error",
                "message": "An internal server error occurred",
                "details": str(e) if os.getenv("DEBUG") else "Please try again later"
            }
        )

# CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:3000", "http://localhost:3001"],  # React dev server
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Data models
class ChatMessage(BaseModel):
    message: str
    chat_history: Optional[List[dict]] = []

class ValidateText(BaseModel):
    validated_text: str
    save_type: Optional[str] = 'auto'

class Document(BaseModel):
    id: str
    name: str
    size: int
    status: str
    created_at: str
    pages: Optional[int] = None
    extracted_text: Optional[str] = None

# Standardized response models
class StandardResponse(BaseModel):
    success: bool
    message: str
    data: Optional[dict] = None
    error: Optional[str] = None
    timestamp: str = datetime.now().isoformat()

class PaginatedResponse(BaseModel):
    success: bool
    data: List[dict]
    pagination: dict
    timestamp: str = datetime.now().isoformat()

class TaskStatusResponse(BaseModel):
    task_id: str
    status: str
    progress: int
    created_at: str
    result: Optional[dict] = None
    error: Optional[str] = None
    message: Optional[str] = None
    current_stage: Optional[str] = None
    redirect_to: Optional[str] = None

# OCR processor initialized without database connection
# Database will be initialized on first use

# Directories
UPLOAD_DIR = Path("uploads")
OUTPUT_DIR = Path("outputs")
UPLOAD_DIR.mkdir(exist_ok=True)
OUTPUT_DIR.mkdir(exist_ok=True)

# Job storage for async OCR tasks (in production, use Redis or database)
jobs = {}
executor = ThreadPoolExecutor(max_workers=3)  # Limit concurrent OCR tasks

# Memory management and cleanup utilities
def cleanup_old_files(directory: Path, max_age_hours: int = 24):
    """Clean up old files to manage disk space"""
    try:
        current_time = time.time()
        for file_path in directory.glob("*"):
            if file_path.is_file():
                file_age = current_time - file_path.stat().st_mtime
                if file_age > (max_age_hours * 3600):
                    try:
                        file_path.unlink()
                        logger.info(f"Cleaned up old file: {file_path.name}")
                    except Exception as e:
                        logger.warning(f"Failed to delete old file {file_path.name}: {e}")
    except Exception as e:
        logger.error(f"Error during file cleanup: {e}")

def cleanup_failed_job_files(task_id: str):
    """Clean up files associated with a failed job"""
    try:
        # Clean up upload files
        for file_path in UPLOAD_DIR.glob(f"*{task_id}*"):
            try:
                file_path.unlink()
                logger.info(f"Cleaned up failed job file: {file_path.name}")
            except Exception as e:
                logger.warning(f"Failed to delete job file {file_path.name}: {e}")
        
        # Clean up output files
        for file_path in OUTPUT_DIR.glob(f"*{task_id}*"):
            try:
                file_path.unlink()
                logger.info(f"Cleaned up failed output file: {file_path.name}")
            except Exception as e:
                logger.warning(f"Failed to delete output file {file_path.name}: {e}")
    except Exception as e:
        logger.error(f"Error during failed job cleanup: {e}")

def check_disk_space(directory: Path, min_free_gb: float = 1.0) -> bool:
    """Check if there's enough disk space"""
    try:
        import shutil
        total, used, free = shutil.disk_usage(directory)
        free_gb = free / (1024**3)
        if free_gb < min_free_gb:
            logger.warning(f"Low disk space: {free_gb:.2f}GB free (minimum: {min_free_gb}GB)")
            return False
        return True
    except Exception as e:
        logger.error(f"Error checking disk space: {e}")
        return True  # Assume OK if check fails

class ProgressTracker:
    """Enhanced progress tracking with detailed stage management"""
    def __init__(self, task_id):
        self.task_id = task_id
        self.stages = {
            'upload': {'min': 0, 'max': 5, 'message': 'File uploaded successfully'},
            'database_save': {'min': 5, 'max': 10, 'message': 'Saving to database...'},
            'analysis': {'min': 10, 'max': 25, 'message': 'Analyzing document structure...'},
            'separation': {'min': 25, 'max': 35, 'message': 'Separating arc diagrams...'},
            'ocr_processing': {'min': 35, 'max': 65, 'message': 'Extracting text from pages...'},
            'vision_processing': {'min': 65, 'max': 85, 'message': 'Processing diagrams...'},
            'combining': {'min': 85, 'max': 95, 'message': 'Combining results...'},
            'finalizing': {'min': 95, 'max': 100, 'message': 'Finalizing processing...'}
        }
    
    def update_progress(self, stage, sub_progress=0.0):
        """Update progress for a specific stage with sub-progress (0-1)"""
        if stage not in self.stages:
            return
            
        stage_info = self.stages[stage]
        progress_range = stage_info['max'] - stage_info['min']
        current_progress = stage_info['min'] + (progress_range * sub_progress)
        
        jobs[self.task_id]['progress'] = min(100, max(0, current_progress))
        jobs[self.task_id]['status_message'] = stage_info['message']
        jobs[self.task_id]['current_stage'] = stage
        
        print(f"[PROGRESS] Task {self.task_id}: {stage} - {current_progress:.1f}% ({stage_info['message']})")

async def save_document_pages_as_images(document_id: str, pdf_path: str, cursor):
    """Save document pages as images for book view display"""
    try:
        from pdf2image import convert_from_path
        import io
        
        # Convert PDF pages to images
        pages = convert_from_path(pdf_path, dpi=150, fmt='JPEG')
        
        for page_num, page_image in enumerate(pages, 1):
            # Convert PIL image to raw JPEG bytes (store raw bytes in DB)
            img_buffer = io.BytesIO()
            page_image.save(img_buffer, format='JPEG', quality=85)
            img_bytes = img_buffer.getvalue()
            
            # Save to database
            cursor.execute("""
                INSERT INTO pages (id, document_id, page_number, page_type, image_data, processing_status, created_at)
                VALUES (%s, %s, %s, %s, %s, %s, %s)
                ON CONFLICT (document_id, page_number) DO UPDATE SET
                    image_data = EXCLUDED.image_data,
                    processing_status = EXCLUDED.processing_status
            """, (
                str(uuid.uuid4()),
                document_id,
                page_num,
                'standard',
                img_bytes,
                'completed',
                datetime.now()
            ))
            
        print(f"[OK] Saved {len(pages)} page images for document {document_id}")
        
    except Exception as e:
        print(f"[WARN] Failed to save document pages as images: {e}")

async def process_fallback_workflow(task_id: str, file_path: str, filename: str, progress_tracker: ProgressTracker, document_id: str):
    """Fallback workflow using Google Vision API for entire document"""
    try:
        progress_tracker.update_progress('vision_processing', 0.0)
        
        if not DiagramTextExtractor:
            raise Exception("DiagramTextExtractor not available")
        diagram_extractor = DiagramTextExtractor()
        loop = asyncio.get_event_loop()
        
        fallback_result = await loop.run_in_executor(
            executor,
            lambda: diagram_extractor.extract_text_from_pdf(
                pdf_path=file_path,
                output_dir=str(OUTPUT_DIR)
            )
        )
        
        progress_tracker.update_progress('vision_processing', 1.0)
        progress_tracker.update_progress('finalizing', 0.5)
        
        # Save fallback results to database
        if ocr_processor and fallback_result:
            try:
                conn = ocr_processor.get_db_connection()
                cursor = conn.cursor()
                
                total_pages = fallback_result.get('total_pages', 0)
                extracted_text = fallback_result.get('combined_text', '')
                
                # Insert or update document record
                cursor.execute("""
                    INSERT INTO documents (id, filename, original_filename, file_size, file_path, processing_status, created_at, total_pages)
                    VALUES (%s, %s, %s, %s, %s, %s, %s, %s)
                    ON CONFLICT (id) DO UPDATE SET
                        processing_status = EXCLUDED.processing_status,
                        total_pages = EXCLUDED.total_pages,
                        updated_at = EXCLUDED.created_at
                """, (document_id, filename, filename, os.path.getsize(file_path), file_path, 'completed', datetime.now(), total_pages))
                
                # Save extracted content
                if extracted_text:
                    cursor.execute("""
                        INSERT INTO extracted_content (id, document_id, content_type, raw_text, processed_text, confidence_score, processing_method, metadata, created_at)
                        VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s)
                        ON CONFLICT (document_id, content_type) DO UPDATE SET
                            raw_text = EXCLUDED.raw_text,
                            processed_text = EXCLUDED.processed_text,
                            metadata = EXCLUDED.metadata,
                            updated_at = CURRENT_TIMESTAMP
                    """, (
                        str(uuid.uuid4()),
                        document_id,
                        'complete_document',
                        extracted_text,
                        extracted_text,
                        0.95,
                        'google_vision_fallback',
                        json.dumps({
                            'workflow_type': 'fallback',
                            'total_pages': total_pages,
                            'processing_date': datetime.now().isoformat()
                        }),
                        datetime.now()
                    ))
                
                # Save page images
                await save_document_pages_as_images(document_id, file_path, cursor)
                
                conn.commit()
                cursor.close()
                conn.close()
                
                progress_tracker.update_progress('finalizing', 1.0)
                
            except Exception as e:
                print(f"[WARN] Failed to save fallback document to database: {e}")
        
        jobs[task_id].update({
            'status': 'completed',
            'progress': 100,
            'status_message': 'Processing completed with fallback workflow',
            'result': {
                'document_id': document_id,
                'workflow_type': 'fallback_google_vision',
                'total_pages': fallback_result.get('total_pages', 0),
                'extracted_text': fallback_result.get('combined_text', ''),
                'pages': fallback_result.get('total_pages', 0),
                'processing_complete': True,
                'message': 'Processed with Google Cloud Vision API (fallback mode)'
            }
        })
        
        print(f"[OK] Fallback workflow completed for task {task_id}")
        
    except Exception as e:
        print(f"[ERROR] Fallback workflow failed: {e}")
        jobs[task_id]['status'] = 'failed'
        jobs[task_id]['error'] = str(e)

async def process_ocr_task_enhanced(task_id: str, file_path: str, filename: str, progress_tracker: ProgressTracker, document_id: str):
    """Enhanced OCR processing with progress tracking"""
    try:
        progress_tracker.update_progress('ocr_processing', 0.0)
        
        if ocr_processor:
            loop = asyncio.get_event_loop()
            if not hasattr(ocr_processor, 'process_pdf_from_frontend'):
                raise AttributeError("OCR processor missing required method")
            result = await loop.run_in_executor(
                executor,
                lambda: getattr(ocr_processor, 'process_pdf_from_frontend')(
                    pdf_path=file_path,
                    original_filename=filename,
                    output_dir=str(OUTPUT_DIR)
                )
            )
            
            progress_tracker.update_progress('ocr_processing', 1.0)
            progress_tracker.update_progress('finalizing', 0.5)
            
            # Update database and finalize
            if result and result.get('status') == 'success':
                if ocr_processor:
                    try:
                        conn = ocr_processor.get_db_connection()
                        cursor = conn.cursor()
                        
                        # Insert or update document
                        cursor.execute("""
                            INSERT INTO documents (id, filename, original_filename, file_size, file_path, processing_status, created_at, total_pages)
                            VALUES (%s, %s, %s, %s, %s, %s, %s, %s)
                            ON CONFLICT (id) DO UPDATE SET
                                processing_status = EXCLUDED.processing_status,
                                total_pages = EXCLUDED.total_pages,
                                updated_at = EXCLUDED.created_at
                        """, (document_id, filename, filename, os.path.getsize(file_path), file_path, 'completed', datetime.now(), result.get('total_pages', 0)))
                        
                        # Save page images
                        await save_document_pages_as_images(document_id, file_path, cursor)
                        
                        conn.commit()
                        cursor.close()
                        conn.close()
                        
                    except Exception as e:
                        print(f"[WARN] Failed to update database: {e}")
                
                progress_tracker.update_progress('finalizing', 1.0)
                
                jobs[task_id].update({
                    'status': 'completed',
                    'progress': 100,
                    'status_message': 'OCR processing completed successfully',
                    'result': {
                        "document_id": document_id,
                        "message": "Document processed successfully",
                        "status": result['status'],
                        "extracted_text": result['extracted_text'],
                        "pages": result['total_pages'],
                        "processing_time_ms": result['processing_time_ms'],
                        "processing_complete": True
                    }
                })
            else:
                raise Exception("OCR processing failed")
        else:
            raise Exception("OCR engine not available")
            
    except Exception as e:
        print(f"[ERROR] Enhanced OCR processing failed: {e}")
        jobs[task_id]['status'] = 'failed'
        jobs[task_id]['error'] = str(e)

# Run periodic cleanup
async def periodic_cleanup():
    """Periodic cleanup task"""
    while True:
        try:
            cleanup_old_files(UPLOAD_DIR, 24)  # Clean uploads older than 24 hours
            cleanup_old_files(OUTPUT_DIR, 72)  # Clean outputs older than 72 hours
            await asyncio.sleep(3600)  # Run every hour
        except Exception as e:
            logger.error(f"Error in periodic cleanup: {e}")
            await asyncio.sleep(3600)

@app.on_event("startup")
async def startup_event():
    """Initialize background tasks on startup"""
    logger.info("Starting background cleanup task")
    asyncio.create_task(periodic_cleanup())
    try:
        # Ensure database schema exists so uploads can be saved
        if ocr_processor and hasattr(ocr_processor, 'ensure_database_initialized'):
            ocr_processor.ensure_database_initialized()
            logger.info("[OK] Database schema ensured on startup")

        # Ensure required unique constraints/indexes exist for UPSERTs
        if ocr_processor:
            conn = ocr_processor.get_db_connection()
            cur = conn.cursor()
            try:
                # Create the unique index used by ON CONFLICT (document_id, page_number)
                cur.execute("""
                    DO $$
                    BEGIN
                        IF NOT EXISTS (
                            SELECT 1 FROM pg_indexes 
                            WHERE schemaname = 'public' AND indexname = 'pages_document_page_unique'
                        ) THEN
                            CREATE UNIQUE INDEX pages_document_page_unique
                            ON pages(document_id, page_number);
                        END IF;
                    END
                    $$;
                """)
                conn.commit()
                logger.info("[OK] Verified unique index on pages(document_id, page_number)")
            finally:
                cur.close()
                conn.close()
    except Exception as e:
        logger.error(f"[WARN] Failed to ensure database on startup: {e}")

async def process_pdf_with_new_workflow(task_id: str, file_path: str, filename: str, document_id: str):
    """
    New workflow: Separate arc diagrams and extract text from remaining content
    """
    try:
        print(f"[WORKFLOW] Starting new PDF workflow for task {task_id}")
        jobs[task_id]['status'] = 'processing'
        jobs[task_id]['progress'] = 10
        
        if not ArcDiagramSeparator or not DiagramTextExtractor:
            raise Exception("Workflow modules not available")
        
        # Step 1: Separate arc diagrams from PDF
        print(f"[PROCESS] Separating arc diagrams from {filename}...")
        jobs[task_id]['progress'] = 25
        arc_separator = ArcDiagramSeparator(similarity_threshold=0.7)
        
        loop = asyncio.get_event_loop()
        separation_result = await loop.run_in_executor(
            executor,
            lambda: arc_separator.separate_arc_diagrams(
                pdf_path=file_path,
                output_dir=str(OUTPUT_DIR)
            )
        )
        jobs[task_id]['progress'] = 50
        
        if separation_result.get('error'):
            print(f"[WARN] Arc diagram separation failed: {separation_result['error']}")
            print(f"[FALLBACK] Falling back to processing all pages with Google Cloud Vision API...")
            jobs[task_id]['progress'] = 60
            
            # Fallback: Process entire PDF with Google Cloud Vision API
            if not DiagramTextExtractor:
                raise Exception("DiagramTextExtractor not available")
            diagram_extractor = DiagramTextExtractor()
            
            fallback_result = await loop.run_in_executor(
                executor,
                lambda: diagram_extractor.extract_text_from_pdf(
                    pdf_path=file_path,
                    output_dir=str(OUTPUT_DIR)
                )
            )
            jobs[task_id]['progress'] = 90
            
            # Save fallback results to database
            if ocr_processor and fallback_result:
                try:
                    # Save document to database
                    conn = ocr_processor.get_db_connection()
                    cursor = conn.cursor()
                    
                    total_pages = fallback_result.get('total_pages', 0)
                    extracted_text = fallback_result.get('combined_text', '')
                    
                    # Insert document record
                    cursor.execute("""
                        INSERT INTO documents (id, filename, original_filename, file_size, file_path, processing_status, created_at, total_pages)
                        VALUES (%s, %s, %s, %s, %s, %s, %s, %s)
                        ON CONFLICT (id) DO UPDATE SET
                            processing_status = EXCLUDED.processing_status,
                            total_pages = EXCLUDED.total_pages
                    """, (document_id, filename, filename, os.path.getsize(file_path), file_path, 'completed', datetime.now(), total_pages))
                    
                    # Convert PDF pages to images and save to database
                    try:
                        import fitz  # PyMuPDF
                    except ImportError:
                        print(f"[ERROR] PyMuPDF not available for task {task_id}")
                        raise Exception("PyMuPDF library not installed")
                        
                    pdf_doc = fitz.open(file_path)
                    
                    for page_num in range(len(pdf_doc)):
                        page = pdf_doc.load_page(page_num)
                        # Use PyMuPDF API with compatibility handling
                        mat = fitz.Matrix(2, 2)  # 2x zoom for better quality
                        try:
                            # Try PyMuPDF API with dynamic method calling for type safety
                            pix = None
                            if hasattr(page, 'get_pixmap'):
                                pix = getattr(page, 'get_pixmap')(matrix=mat)
                            elif hasattr(page, 'getPixmap'):
                                pix = getattr(page, 'getPixmap')(matrix=mat)
                            else:
                                # Try without matrix parameter as fallback
                                if hasattr(page, 'get_pixmap'):
                                    pix = getattr(page, 'get_pixmap')()
                                elif hasattr(page, 'getPixmap'):
                                    pix = getattr(page, 'getPixmap')()
                                else:
                                    raise AttributeError("No pixmap method available")
                                    
                            if pix is None:
                                print(f"Failed to create pixmap for page {page_num + 1}")
                                continue
                        except Exception as pix_error:
                            print(f"PyMuPDF pixmap error on page {page_num + 1}: {pix_error}")
                            # Skip this page if we can't render it
                            continue
                        img_data = pix.tobytes("jpeg")
                        
                        # Save page to database
                        cursor.execute("""
                            INSERT INTO pages (document_id, page_number, image_data)
                            VALUES (%s, %s, %s)
                            ON CONFLICT (document_id, page_number) DO UPDATE SET
                                image_data = EXCLUDED.image_data
                        """, (document_id, page_num + 1, img_data))
                    
                    pdf_doc.close()
                    
                    # Save extracted content
                    cursor.execute("""
                        INSERT INTO extracted_content (document_id, raw_text, processed_text, metadata)
                        VALUES (%s, %s, %s, %s)
                        ON CONFLICT (document_id) DO UPDATE SET
                            raw_text = EXCLUDED.raw_text,
                            processed_text = EXCLUDED.processed_text,
                            metadata = EXCLUDED.metadata
                    """, (document_id, extracted_text, extracted_text, 
                          json.dumps({"workflow": "fallback_google_vision", "total_pages": total_pages})))
                    
                    conn.commit()
                    cursor.close()
                    conn.close()
                    
                    print(f"[OK] Fallback document and pages saved to database for task {task_id}")
                    
                except Exception as e:
                    print(f"[WARN] Failed to save fallback document to database: {e}")
            
            jobs[task_id].update({
                'status': 'completed',
                'result': {
                    'document_id': document_id,
                    'workflow_type': 'fallback_google_vision',
                    'total_pages': fallback_result.get('total_pages', 0),
                    'extracted_text': fallback_result.get('combined_text', ''),
                    'pages': fallback_result.get('total_pages', 0),  # Return pages count, not pages array
                    'message': 'Processed with Google Cloud Vision API (fallback mode)'
                }
            })
            
            print(f"[OK] Fallback workflow completed for task {task_id}")
            return
        
        # Step 2: Process non-arc pages with existing OCR engine if available
        text_extraction_result = {"pages": [], "combined_text": "", "total_pages": 0}
        
        if separation_result['non_arc_pages_count'] > 0 and ocr_processor:
            print(f"[PROCESS] Processing {separation_result['non_arc_pages_count']} non-arc pages with OCR engine...")
            
            # Use the non-arc PDF if it was created
            non_arc_pdf_path = separation_result.get('non_arc_pdf_path')
            if non_arc_pdf_path and Path(non_arc_pdf_path).exists():
                try:
                    if not ocr_processor:
                        raise Exception("OCR processor not available")
                        
                    if not hasattr(ocr_processor, 'process_pdf_from_frontend'):
                        raise AttributeError("OCR processor missing required method")
                        
                    ocr_result = await loop.run_in_executor(
                        executor,
                        lambda: getattr(ocr_processor, 'process_pdf_from_frontend')(
                            pdf_path=non_arc_pdf_path,
                            original_filename=f"non_arc_{filename}",
                            output_dir=str(OUTPUT_DIR)
                        )
                    )
                    
                    text_extraction_result = {
                        "pages": [{"page_number": i+1, "text": ocr_result.get('extracted_text', '')} for i in range(ocr_result.get('total_pages', 0))],
                        "combined_text": ocr_result.get('extracted_text', ''),
                        "total_pages": ocr_result.get('total_pages', 0)
                    }
                except Exception as e:
                    print(f"[WARN] OCR processing failed: {e}")
                    text_extraction_result = {"pages": [], "combined_text": "", "total_pages": 0}
        
        # Step 3: Extract text from arc diagrams using Google Vision API
        arc_text_result = {"pages": [], "combined_text": "", "total_pages": 0}
        
        if separation_result['arc_pages_count'] > 0:
            print(f"[EXTRACT] Extracting text from {separation_result['arc_pages_count']} arc diagram pages using Google Vision API...")
            
            arc_pdf_path = separation_result.get('arc_pdf_path')
            if arc_pdf_path and Path(arc_pdf_path).exists():
                text_extractor = DiagramTextExtractor()
                arc_text_result = await loop.run_in_executor(
                    executor,
                    lambda: text_extractor.extract_text_from_pdf(
                        pdf_path=arc_pdf_path,
                        output_dir=str(OUTPUT_DIR / "extracted_images")
                    )
                )
        
        # Step 4: Combine results
        combined_text = ""
        total_pages = separation_result['total_pages']
        
        if arc_text_result['combined_text']:
            combined_text += f"\n\n=== ARC DIAGRAMS ({arc_text_result['total_pages']} pages) ===\n"
            combined_text += arc_text_result['combined_text']
        
        if text_extraction_result['combined_text']:
            combined_text += f"\n\n=== TEXT CONTENT ({text_extraction_result['total_pages']} pages) ===\n"
            combined_text += text_extraction_result['combined_text']
        
        # Step 5: Save document and pages to database for PDF preview
        if ocr_processor:
            try:
                # Save document to database
                conn = ocr_processor.get_db_connection()
                cursor = conn.cursor()
                
                # Insert document record
                cursor.execute("""
                    INSERT INTO documents (id, filename, original_filename, file_size, file_path, processing_status, created_at, total_pages)
                    VALUES (%s, %s, %s, %s, %s, %s, %s, %s)
                    ON CONFLICT (id) DO UPDATE SET
                        processing_status = EXCLUDED.processing_status,
                        total_pages = EXCLUDED.total_pages
                """, (document_id, filename, filename, os.path.getsize(file_path), file_path, 'completed', datetime.now(), total_pages))
                
                # Convert PDF pages to images and save to database
                try:
                    import fitz  # PyMuPDF
                except ImportError:
                    print(f"[ERROR] PyMuPDF not available for task {task_id}")
                    raise Exception("PyMuPDF library not installed")
                    
                pdf_doc = fitz.open(file_path)
                
                for page_num in range(len(pdf_doc)):
                    page = pdf_doc.load_page(page_num)
                    # Use PyMuPDF API with dynamic method calling for type safety
                    mat = fitz.Matrix(2, 2)  # 2x zoom for better quality
                    try:
                        # Try different PyMuPDF API methods based on version using getattr
                        pix = None
                        if hasattr(page, 'get_pixmap'):
                            # Newer PyMuPDF API
                            pix = getattr(page, 'get_pixmap')(matrix=mat)
                        elif hasattr(page, 'getPixmap'):
                            # Older PyMuPDF API
                            pix = getattr(page, 'getPixmap')(matrix=mat)
                        else:
                            # Try without matrix parameter as fallback
                            if hasattr(page, 'get_pixmap'):
                                pix = getattr(page, 'get_pixmap')()
                            elif hasattr(page, 'getPixmap'):
                                pix = getattr(page, 'getPixmap')()
                            else:
                                raise AttributeError("No pixmap method available")
                                
                        if pix is None:
                            print(f"Failed to create pixmap for page {page_num + 1}")
                            continue
                            
                    except Exception as pix_error:
                        print(f"PyMuPDF pixmap error on page {page_num + 1}: {pix_error}")
                        # Skip this page if we can't render it
                        continue
                        
                    img_data = pix.tobytes("jpeg")
                    
                    # Save page to database
                    cursor.execute("""
                        INSERT INTO pages (document_id, page_number, image_data)
                        VALUES (%s, %s, %s)
                        ON CONFLICT (document_id, page_number) DO UPDATE SET
                            image_data = EXCLUDED.image_data
                    """, (document_id, page_num + 1, img_data))
                
                pdf_doc.close()
                
                # Save extracted content
                cursor.execute("""
                    INSERT INTO extracted_content (id, document_id, content_type, raw_text, processed_text, metadata, created_at)
                    VALUES (%s, %s, %s, %s, %s, %s, %s)
                    ON CONFLICT (document_id, content_type) DO UPDATE SET
                        raw_text = EXCLUDED.raw_text,
                        processed_text = EXCLUDED.processed_text,
                        metadata = EXCLUDED.metadata,
                        updated_at = CURRENT_TIMESTAMP
                """, (str(uuid.uuid4()), document_id, 'complete_document', combined_text.strip(), combined_text.strip(), 
                      json.dumps({"workflow": "new_workflow", "total_pages": total_pages}), datetime.now()))
                
                # Update document status to completed
                cursor.execute("""
                    UPDATE documents 
                    SET processing_status = 'completed', updated_at = CURRENT_TIMESTAMP
                    WHERE id = %s
                """, (document_id,))
                
                conn.commit()
                cursor.close()
                conn.close()
                
                print(f"[OK] Document and pages saved to database for task {task_id}")
                
            except Exception as e:
                print(f"[WARN] Failed to save document to database: {e}")
        
        # Calculate processing time (mock)
        processing_time = 2000  # milliseconds
        
        jobs[task_id]['status'] = 'completed'
        jobs[task_id]['result'] = {
            "document_id": document_id,
            "message": "Document processed with new workflow",
            "status": "completed",
            "extracted_text": combined_text.strip(),
            "pages": total_pages,  # Ensure this is always a number
            "processing_time_ms": processing_time,
            "workflow_details": {
                "arc_pages": separation_result['arc_pages_count'],
                "non_arc_pages": separation_result['non_arc_pages_count'],
                "total_pages": total_pages,
                "arc_pdf_path": separation_result.get('arc_pdf_path'),
                "non_arc_pdf_path": separation_result.get('non_arc_pdf_path')
            }
        }
        
        print(f"[OK] New workflow completed for task {task_id}")
        
    except Exception as e:
        print(f"[ERROR] New workflow failed for task {task_id}: {e}")
        jobs[task_id]['status'] = 'failed'
        jobs[task_id]['error'] = str(e)
        # Clean up file if processing failed
        if Path(file_path).exists():
            Path(file_path).unlink()

@app.get("/")
async def root():
    return {"message": "OCR AI Assistant API", "status": "running"}

@app.get("/health")
async def health_check():
    """Comprehensive health check with dependency status"""
    try:
        health_status = {
            "status": "healthy",
            "timestamp": datetime.now().isoformat(),
            "service": "OCR AI Assistant API",
            "version": "1.0.0",
            "dependencies": {}
        }
        
        # Check database
        try:
            if ocr_processor:
                conn = ocr_processor.get_db_connection()
                cursor = conn.cursor()
                cursor.execute("SELECT 1")
                cursor.fetchone()
                cursor.close()
                conn.close()
                health_status["dependencies"]["database"] = "healthy"
            else:
                health_status["dependencies"]["database"] = "unavailable"
        except Exception as e:
            health_status["dependencies"]["database"] = f"error: {str(e)}"
            health_status["status"] = "degraded"
        
        # Check OCR engine
        if ocr_processor:
            health_status["dependencies"]["ocr_engine"] = "available"
        else:
            health_status["dependencies"]["ocr_engine"] = "unavailable"
            health_status["status"] = "degraded"
        
        # Check workflow modules
        if ArcDiagramSeparator and DiagramTextExtractor:
            health_status["dependencies"]["workflow_modules"] = "available"
        else:
            health_status["dependencies"]["workflow_modules"] = "limited"
        
        # Check disk space
        try:
            import shutil
            total, used, free = shutil.disk_usage(UPLOAD_DIR)
            free_gb = free / (1024**3)
            health_status["dependencies"]["disk_space"] = f"{free_gb:.2f}GB free"
            if free_gb < 1.0:
                health_status["status"] = "degraded"
        except Exception as e:
            health_status["dependencies"]["disk_space"] = f"error: {str(e)}"
        
        return health_status
        
    except Exception as e:
        logger.error(f"Health check failed: {e}")
        return {
            "status": "error",
            "error": str(e),
            "timestamp": datetime.now().isoformat()
        }

# Removed /health/db endpoint to avoid database connection conflicts

@app.get("/test-sentry")
async def test_sentry():
    """Test endpoint to verify Sentry error reporting"""
    try:
        # Intentionally trigger an error for testing
        raise Exception("This is a test error for Sentry integration")
    except Exception as e:
        # Re-raise to let Sentry capture it
        raise HTTPException(status_code=500, detail="Sentry test error triggered successfully")

async def process_ocr_task(task_id: str, file_path: str, filename: str):
    """Background OCR processing function"""
    try:
        print(f"[OCR] Starting OCR processing for task {task_id}")
        jobs[task_id]['status'] = 'processing'
        jobs[task_id]['progress'] = 20
        
        if ocr_processor:
            # Run OCR processing in thread pool to avoid blocking
            jobs[task_id]['progress'] = 40
            loop = asyncio.get_event_loop()
            try:
                if not hasattr(ocr_processor, 'process_pdf_from_frontend'):
                    raise AttributeError("OCR processor missing required method")
                    
                result = await loop.run_in_executor(
                    executor,
                    lambda: getattr(ocr_processor, 'process_pdf_from_frontend')(
                        pdf_path=file_path,
                        original_filename=filename,
                        output_dir=str(OUTPUT_DIR)
                    )
                )
                
                jobs[task_id]['progress'] = 90
                jobs[task_id]['status'] = 'completed'
                jobs[task_id]['result'] = {
                    "document_id": result['document_id'],
                    "message": "Document processed successfully",
                    "status": result['status'],
                    "extracted_text": result['extracted_text'],
                    "pages": result['total_pages'],
                    "processing_time_ms": result['processing_time_ms']
                }
                print(f"[OK] OCR processing completed for task {task_id}")
            except Exception as e:
                print(f"[ERROR] OCR processing failed for task {task_id}: {e}")
                jobs[task_id]['status'] = 'failed'
                jobs[task_id]['error'] = f'OCR processing failed: {str(e)}'
        else:
            jobs[task_id]['status'] = 'failed'
            jobs[task_id]['error'] = 'OCR engine not available'
            
    except Exception as e:
        print(f"[ERROR] OCR processing failed for task {task_id}: {e}")
        jobs[task_id]['status'] = 'failed'
        jobs[task_id]['error'] = str(e)
        # Clean up file if processing failed
        if Path(file_path).exists():
            Path(file_path).unlink()

@app.post("/upload/")
async def upload_document(file: UploadFile = File(...)):
    """Upload document and start async OCR processing with improved memory management"""
    
    # Validate file type
    if not file.filename or not file.filename.lower().endswith('.pdf'):
        raise HTTPException(status_code=400, detail="Only PDF files are allowed")
    
    # Check file size (50MB limit)
    max_size = 50 * 1024 * 1024  # 50MB
    if file.size and file.size > max_size:
        raise HTTPException(status_code=400, detail="File size must be less than 50MB")
    
    # Check disk space
    if not check_disk_space(UPLOAD_DIR, 1.0):  # Require at least 1GB free
        raise HTTPException(status_code=507, detail="Insufficient disk space")
    
    # Generate unique task ID
    task_id = f"task_{int(datetime.now().timestamp() * 1000)}"
    doc_id = str(uuid.uuid4())
    
    # Save uploaded file with better error handling
    file_path = UPLOAD_DIR / f"{doc_id}_{file.filename}"
    
    try:
        # Stream file to disk to handle large files efficiently
        with open(file_path, "wb") as buffer:
            chunk_size = 8192  # 8KB chunks
            while chunk := await file.read(chunk_size):
                buffer.write(chunk)
        
        # Verify file was written correctly
        if not file_path.exists() or file_path.stat().st_size == 0:
            raise Exception("File upload verification failed")
        
        logger.info(f"File uploaded successfully: {file.filename} ({file_path.stat().st_size} bytes)")
        
        # Get filename with fallback
        filename = file.filename or "unknown.pdf"
        
        # Initialize job status with proper tracking
        jobs[task_id] = {
            'status': 'uploaded',
            'result': None,
            'error': None,
            'progress': 5,  # Start at 5% after upload
            'status_message': 'File uploaded successfully',
            'current_stage': 'upload',
            'created_at': datetime.now().isoformat(),
            'file_info': {
                'filename': filename,
                'size': file_path.stat().st_size,
                'doc_id': doc_id
            }
        }
        
        # Create initial document record so it appears in Documents list immediately
        if ocr_processor:
            try:
                conn = ocr_processor.get_db_connection()
                cursor = conn.cursor()
                cursor.execute("""
                    INSERT INTO documents (id, filename, original_filename, file_size, file_path, processing_status, created_at)
                    VALUES (%s, %s, %s, %s, %s, %s, %s)
                    ON CONFLICT (id) DO NOTHING
                """, (doc_id, filename, filename, file_path.stat().st_size, str(file_path), 'processing', datetime.now()))
                conn.commit()
                cursor.close()
                conn.close()
                logger.info(f"Initial document record created for {doc_id}")
            except Exception as e:
                logger.warning(f"Failed to create initial document record: {e}")
        
        # Start new PDF workflow processing in background
        asyncio.create_task(process_pdf_with_new_workflow(task_id, str(file_path), filename, doc_id))
        
        # Immediately respond with task ID and document ID
        return {
            "task_id": task_id,
            "document_id": doc_id,
            "message": "Document uploaded successfully. Processing started.",
            "status": "uploaded",
            "file_size": file_path.stat().st_size
        }
        
    except Exception as e:
        # Clean up file if upload failed
        if file_path.exists():
            try:
                file_path.unlink()
            except Exception as cleanup_error:
                logger.error(f"Failed to cleanup failed upload: {cleanup_error}")
        
        logger.error(f"Upload failed for {file.filename}: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to upload document: {str(e)}")

@app.get("/task-status/{task_id}", response_model=TaskStatusResponse)
async def get_task_status(task_id: str):
    """Check the status of an OCR processing task with enhanced progress information"""
    
    if task_id not in jobs:
        raise HTTPException(status_code=404, detail="Task not found")
    
    job = jobs[task_id]
    
    response_data = {
        "task_id": task_id,
        "status": job['status'],
        "created_at": job['created_at'],
        "progress": job.get('progress', 0)
    }
    
    # Add enhanced progress information
    if 'status_message' in job:
        response_data['message'] = job['status_message']
    elif job['status'] in ['uploaded', 'processing']:
        response_data['message'] = f"Task is {job['status']}..."
    
    # Add current stage information
    if 'current_stage' in job:
        response_data['current_stage'] = job['current_stage']
    
    if job['status'] == 'completed' and job.get('result'):
        response_data['result'] = job['result']
        response_data['progress'] = 100
        # Add redirect information for frontend navigation
        if 'document_id' in job['result']:
            response_data['redirect_to'] = f"/editor/{job['result']['document_id']}"
    elif job['status'] == 'failed' and job.get('error'):
        response_data['error'] = job['error']
    
    return response_data

def simulate_ocr_processing(file_path: Path) -> str:
    """Simulate OCR processing - replace with actual OCR engine integration"""
    
    # This is where you would integrate with the actual OCR engine
    # For now, return a sample extracted text
    
    sample_text = f"""LEASE AGREEMENT

This Lease Agreement ("Agreement") is entered into on [DATE] between [LANDLORD NAME] ("Landlord") and [TENANT NAME] ("Tenant").

PROPERTY DETAILS:
Address: [PROPERTY ADDRESS]
Unit: [UNIT NUMBER]
City, State, ZIP: [CITY, STATE ZIP]

LEASE TERMS:
• Lease Term: [START DATE] to [END DATE]
• Monthly Rent: $[AMOUNT]
• Security Deposit: $[AMOUNT]
• Pet Deposit: $[AMOUNT] (if applicable)

TENANT RESPONSIBILITIES:
1. Pay rent on time each month
2. Maintain property in good condition
3. Follow all building rules and regulations
4. Provide 30 days notice before moving out

LANDLORD RESPONSIBILITIES:
1. Maintain structural integrity of property
2. Ensure all utilities are functioning
3. Respond to maintenance requests promptly
4. Respect tenant's right to quiet enjoyment

ADDITIONAL TERMS:
• No smoking allowed on premises
• Maximum occupancy: [NUMBER] persons
• Parking: [PARKING DETAILS]
• Utilities included: [UTILITIES LIST]

This agreement is binding upon both parties and their successors.

Landlord Signature: _________________________ Date: _________

Tenant Signature: _________________________ Date: _________

[Document processed from: {file_path.name}]
"""
    
    return sample_text

@app.get("/documents/mock")
async def get_documents_mock():
    """Mock documents endpoint for testing"""
    print("[MOCK] GET /documents/mock - Returning mock data")
    return [
        {
            "id": "mock-1",
            "name": "Sample Document 1.pdf",
            "filename": "sample1.pdf",
            "size": 1024000,
            "status": "completed",
            "pages": 5,
            "created_at": "2024-01-01T10:00:00Z"
        },
        {
            "id": "mock-2",
            "name": "Sample Document 2.pdf",
            "filename": "sample2.pdf",
            "size": 2048000,
            "status": "processing",
            "pages": 3,
            "created_at": "2024-01-01T11:00:00Z"
        }
    ]

@app.get("/documents/", response_model=PaginatedResponse)
async def get_documents(page: int = 1, limit: int = 20, status: Optional[str] = None):
    """Get documents with pagination and optional status filtering"""
    try:
        logger.info(f"[API] GET /documents/ - Fetching documents (page={page}, limit={limit}, status={status})")
        
        if not ocr_processor:
            logger.error("[ERROR] OCR processor not available")
            raise HTTPException(status_code=500, detail="OCR processor not available")
        
        # Validate pagination parameters
        if page < 1:
            page = 1
        if limit < 1 or limit > 100:
            limit = 20
        
        offset = (page - 1) * limit
        
        # Add timeout handling for database operations
        import asyncio
        import concurrent.futures
        
        def get_documents_sync():
            conn = None
            cursor = None
            try:
                if not ocr_processor:
                    raise Exception("OCR processor not available")
                    
                conn = ocr_processor.get_db_connection()
                cursor = conn.cursor()
                
                # Set query timeout
                cursor.execute("SET statement_timeout = '15s'")
                
                # Build WHERE clause for status filtering
                where_clause = ""
                params = []
                if status:
                    where_clause = "WHERE processing_status = %s"
                    params.append(status)
                
                # Get total count
                count_query = f"SELECT COUNT(*) FROM documents {where_clause}"
                cursor.execute(count_query, params)
                total_count = cursor.fetchone()[0]
                
                # Get paginated documents
                query = f"""
                    SELECT id,
                           filename,
                           COALESCE(original_filename, filename) AS original_filename,
                           file_size,
                           processing_status AS status,
                           total_pages AS pages,
                           COALESCE(upload_date, created_at) AS created_at
                    FROM documents
                    {where_clause}
                    ORDER BY COALESCE(upload_date, created_at) DESC NULLS LAST
                    LIMIT %s OFFSET %s
                """
                
                cursor.execute(query, params + [limit, offset])
                rows = cursor.fetchall()
                
                logger.info(f"[INFO] Found {len(rows)} documents (total: {total_count})")
                
                documents = []
                for row in rows:
                    documents.append({
                        "id": str(row[0]),
                        "name": row[2],  # original_filename
                        "filename": row[1],
                        "size": row[3],
                        "status": row[4],
                        "pages": row[5],
                        "created_at": row[6].isoformat() if row[6] else None
                    })
                
                # Calculate pagination info
                total_pages = (total_count + limit - 1) // limit
                has_next = page < total_pages
                has_prev = page > 1
                
                pagination = {
                    "current_page": page,
                    "per_page": limit,
                    "total_items": total_count,
                    "total_pages": total_pages,
                    "has_next": has_next,
                    "has_prev": has_prev
                }
                
                return {
                    "success": True,
                    "data": documents,
                    "pagination": pagination
                }
                
            finally:
                if cursor:
                    cursor.close()
                if conn:
                    conn.close()
        
        # Execute with timeout
        loop = asyncio.get_event_loop()
        with concurrent.futures.ThreadPoolExecutor() as executor:
            try:
                result = await asyncio.wait_for(
                    loop.run_in_executor(executor, get_documents_sync),
                    timeout=20.0  # 20 second timeout
                )
                logger.info(f"[OK] Successfully returning {len(result['data'])} documents")
                return result
            except asyncio.TimeoutError:
                logger.error("⏰ Database query timed out")
                raise HTTPException(status_code=504, detail="Database query timeout - please try again")
            except Exception as db_error:
                logger.error(f"💥 Database error: {db_error}")
                # Return structured fallback data
                return {
                    "success": True,
                    "data": [
                        {
                            "id": "fallback-1",
                            "name": "Fallback Document.pdf",
                            "filename": "fallback.pdf",
                            "size": 512000,
                            "status": "completed",
                            "pages": 1,
                            "created_at": "2024-01-01T12:00:00Z"
                        }
                    ],
                    "pagination": {
                        "current_page": 1,
                        "per_page": limit,
                        "total_items": 1,
                        "total_pages": 1,
                        "has_next": False,
                        "has_prev": False
                    }
                }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"[ERROR] Error retrieving documents: {e}")
        logger.error(f"Traceback: {traceback.format_exc()}")
        raise HTTPException(status_code=500, detail=f"Failed to retrieve documents: {str(e)}")

@app.get("/documents/{doc_id}/status")
async def get_document_status(doc_id: str):
    """Get document processing status from database"""
    try:
        if not ocr_processor:
            raise HTTPException(status_code=500, detail="OCR processor not available")
            
        conn = ocr_processor.get_db_connection()
        cursor = conn.cursor()
        
        cursor.execute("""
            SELECT processing_status, total_pages
            FROM documents 
            WHERE id = %s
        """, (doc_id,))
        
        result = cursor.fetchone()
        cursor.close()
        conn.close()
        
        if not result:
            raise HTTPException(status_code=404, detail="Document not found")
        
        status, total_pages = result
        progress = 100 if status == "completed" else 50 if status == "processing" else 0
        
        return {
            "id": doc_id,
            "status": status,
            "progress": progress,
            "total_pages": total_pages
        }
        
    except HTTPException:
        raise
    except Exception as e:
        print(f"Error getting document status: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to get document status: {str(e)}")

@app.get("/documents/{doc_id}/content")
async def get_document_content(doc_id: str):
    """Get document content from database"""
    try:
        if not ocr_processor:
            raise HTTPException(status_code=500, detail="OCR processor not available")
            
        # Get complete text from extracted_content and document info
        conn = ocr_processor.get_db_connection()
        cursor = conn.cursor()
        
        # First get document info
        cursor.execute("""
            SELECT total_pages 
            FROM documents 
            WHERE id = %s
        """, (doc_id,))
        
        doc_result = cursor.fetchone()
        if not doc_result:
            cursor.close()
            conn.close()
            raise HTTPException(status_code=404, detail="Document not found")
        
        total_pages = doc_result[0]
        
        # Prefer validated text if present, then complete_document
        complete_text = ""

        # 1) Validated text (highest priority)
        cursor.execute("""
            SELECT processed_text 
            FROM extracted_content 
            WHERE document_id = %s AND content_type = 'validated_document'
            ORDER BY updated_at DESC NULLS LAST, created_at DESC
            LIMIT 1
        """, (doc_id,))
        row = cursor.fetchone()
        if row and row[0]:
            complete_text = row[0]
        else:
            # 2) Complete document text
            cursor.execute("""
                SELECT raw_text 
                FROM extracted_content 
                WHERE document_id = %s AND content_type = 'complete_document'
                ORDER BY created_at DESC
                LIMIT 1
            """, (doc_id,))
            result = cursor.fetchone()
            complete_text = result[0] if result else ""
        
        # If no complete document found, try to get all page content
        if not complete_text:
            cursor.execute("""
                SELECT raw_text 
                FROM extracted_content 
                WHERE document_id = %s AND page_id IS NOT NULL
                ORDER BY metadata->>'page_number'
            """, (doc_id,))
            
            page_results = cursor.fetchall()
            if page_results:
                complete_text = "\n\n".join([row[0] for row in page_results if row[0]])
        
        cursor.close()
        conn.close()
        
        return {
            "id": doc_id,
            "text": complete_text,
            "pages": total_pages
        }
        
    except HTTPException:
        raise
    except Exception as e:
        print(f"Error getting document content: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to get document content: {str(e)}")

@app.post("/validate/{doc_id}")
async def validate_document(doc_id: str, data: ValidateText):
    """Update validated text for a document in database"""
    try:
        if not ocr_processor:
            raise HTTPException(status_code=500, detail="OCR processor not available")
            
        conn = ocr_processor.get_db_connection()
        cursor = conn.cursor()
        
        # Check if document exists
        cursor.execute("SELECT id FROM documents WHERE id = %s", (doc_id,))
        if not cursor.fetchone():
            cursor.close()
            conn.close()
            raise HTTPException(status_code=404, detail="Document not found")
        
        # Check if validated content already exists
        cursor.execute("""
            SELECT id FROM extracted_content 
            WHERE document_id = %s AND content_type = %s
        """, (doc_id, 'validated_document'))
        
        existing_record = cursor.fetchone()
        
        if existing_record:
            # Update existing validated text
            cursor.execute("""
                UPDATE extracted_content 
                SET processed_text = %s, raw_text = %s, updated_at = CURRENT_TIMESTAMP,
                    metadata = %s
                WHERE document_id = %s AND content_type = %s
            """, (
                data.validated_text,
                data.validated_text,
                json.dumps({'validated_by': 'user', 'validation_date': datetime.now().isoformat()}),
                doc_id,
                'validated_document'
            ))
        else:
            # Insert new validated text
            cursor.execute("""
                INSERT INTO extracted_content 
                (document_id, content_type, raw_text, processed_text, metadata)
                VALUES (%s, %s, %s, %s, %s)
            """, (
                doc_id, 
                'validated_document', 
                data.validated_text, 
                data.validated_text,
                json.dumps({'validated_by': 'user', 'validation_date': datetime.now().isoformat()})
            ))
        
        # Update document timestamp
        cursor.execute("""
            UPDATE documents 
            SET updated_at = CURRENT_TIMESTAMP 
            WHERE id = %s
        """, (doc_id,))
        
        conn.commit()
        cursor.close()
        conn.close()
        
        return {"message": "Document updated successfully"}
        
    except HTTPException:
        raise
    except Exception as e:
        print(f"Error validating document: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to update document: {str(e)}")

@app.post("/documents/{doc_id}/auto-save")
async def auto_save_document(
    doc_id: str, 
    data: ValidateText
):
    """Auto-save document content with version tracking"""
    try:
        if not ocr_processor:
            raise HTTPException(status_code=500, detail="OCR processor not available")
        
        # Parse save_type from request body if provided
        save_type = getattr(data, 'save_type', 'auto')
        
        conn = ocr_processor.get_db_connection()
        cursor = conn.cursor()
        
        # Check if document exists
        cursor.execute("SELECT id FROM documents WHERE id = %s", (doc_id,))
        if not cursor.fetchone():
            # Gracefully create a minimal document record to avoid 404 on first auto-save
            try:
                cursor.execute(
                    """
                    INSERT INTO documents (id, filename, original_filename, file_size, processing_status, created_at)
                    VALUES (%s, %s, %s, %s, %s, %s)
                    ON CONFLICT (id) DO NOTHING
                    """,
                    (doc_id, 'unknown.pdf', 'unknown.pdf', 1, 'completed', datetime.now())
                )
            except Exception as insert_err:
                # If creation fails, return a clear error
                conn.rollback()
                cursor.close()
                conn.close()
                raise HTTPException(status_code=404, detail=f"Document not found and could not be created: {insert_err}")
        
        # Update or insert validated content
        cursor.execute("""
            INSERT INTO extracted_content (id, document_id, content_type, raw_text, processed_text, confidence_score, processing_method, metadata, created_at)
            VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s)
            ON CONFLICT (document_id, content_type) DO UPDATE SET
                processed_text = EXCLUDED.processed_text,
                raw_text = EXCLUDED.raw_text,
                metadata = EXCLUDED.metadata,
                updated_at = CURRENT_TIMESTAMP
        """, (
            str(uuid.uuid4()),
            doc_id,
            'validated_document',
            data.validated_text,
            data.validated_text,
            1.0,  # Full confidence for user-validated text
            f'{save_type}_save',
            json.dumps({
                'save_type': save_type,
                'saved_by': 'user',
                'save_date': datetime.now().isoformat()
            }),
            datetime.now()
        ))
        
        # Update document timestamp
        cursor.execute("""
            UPDATE documents 
            SET updated_at = CURRENT_TIMESTAMP
            WHERE id = %s
        """, (doc_id,))
        
        conn.commit()
        cursor.close()
        conn.close()
        
        return {
            "success": True,
            "message": f"Document {save_type}-saved successfully",
            "saved_at": datetime.now().isoformat(),
            "save_type": save_type
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Auto-save failed for document {doc_id}: {e}")
        raise HTTPException(status_code=500, detail="Auto-save failed")

@app.post("/chat/all")
async def chat_with_all_documents(chat_data: ChatMessage):
    """Chat with AI about all documents in the database"""
    try:
        if not ocr_processor:
            raise HTTPException(status_code=500, detail="OCR processor not available")
            
        conn = ocr_processor.get_db_connection()
        cursor = conn.cursor()
        
        # Get all documents and their content
        cursor.execute("""
            SELECT d.id, d.original_filename, ec.raw_text
            FROM documents d
            LEFT JOIN extracted_content ec ON d.id = ec.document_id 
                AND ec.metadata->>'content_type' = 'complete_document'
            WHERE d.processing_status = 'completed'
            ORDER BY d.upload_date DESC
        """)
        
        documents = cursor.fetchall()
        
        # Combine all document texts
        all_documents_text = ""
        document_summaries = []
        documents_with_content = 0
        
        for doc_id, filename, content in documents:
            document_summaries.append(f"- {filename}")
            if content:
                all_documents_text += f"\n\n=== Document: {filename} ===\n{content}"
                documents_with_content += 1
        
        # If we have documents but no content, provide a different message
        if not all_documents_text and documents:
            all_documents_text = f"Found {len(documents)} documents but no extracted content available."
        elif not all_documents_text:
            all_documents_text = "No processed documents found in the library."
        
        # Generate AI response based on all documents
        ai_response = generate_ai_response_all_docs(chat_data.message, all_documents_text)
        
        # Add document list to response if user asks about documents
        if "documents" in chat_data.message.lower() or "files" in chat_data.message.lower():
            if document_summaries:
                ai_response += f"\n\nYour current document library includes:\n" + "\n".join(document_summaries)
            else:
                ai_response += "\n\nYour document library is currently empty. Please upload some documents to get started."
        
        cursor.close()
        conn.close()
        
        return {
            "response": ai_response,
            "documents_count": len(documents),
            "chat_history": (chat_data.chat_history or []) + [
                {"role": "user", "content": chat_data.message, "timestamp": datetime.now().isoformat()},
                {"role": "assistant", "content": ai_response, "timestamp": datetime.now().isoformat()}
            ]
        }
        
    except HTTPException:
        raise
    except Exception as e:
        print(f"Error in global chat: {e}")
        raise HTTPException(status_code=500, detail=f"Chat failed: {str(e)}")

@app.post("/chat/{doc_id}")
async def chat_with_document(doc_id: str, chat_data: ChatMessage):
    """Chat with AI about the document using database"""
    try:
        if not ocr_processor:
            raise HTTPException(status_code=500, detail="OCR processor not available")
            
        conn = ocr_processor.get_db_connection()
        cursor = conn.cursor()
        
        # Check if document exists and get content
        cursor.execute("""
            SELECT d.id, ec.raw_text
            FROM documents d
            LEFT JOIN extracted_content ec ON d.id = ec.document_id 
                AND ec.content_type = 'complete_document'
            WHERE d.id = %s
        """, (doc_id,))
        
        result = cursor.fetchone()
        if not result:
            cursor.close()
            conn.close()
            raise HTTPException(status_code=404, detail="Document not found")
        
        document_text = result[1] or ""
        
        # Get or create chat session
        cursor.execute("""
            SELECT id FROM chat_sessions WHERE document_id = %s
        """, (doc_id,))
        
        existing_session = cursor.fetchone()
        
        if existing_session:
            session_id = existing_session[0]
            # Update timestamp
            cursor.execute("""
                UPDATE chat_sessions SET updated_at = CURRENT_TIMESTAMP WHERE id = %s
            """, (session_id,))
        else:
            # Create new session
            cursor.execute("""
                INSERT INTO chat_sessions (document_id, session_name)
                VALUES (%s, %s)
                RETURNING id
            """, (doc_id, f"Chat session for document {doc_id}"))
            session_id = cursor.fetchone()[0]
        
        # Add user message to database
        cursor.execute("""
            INSERT INTO chat_messages (session_id, document_id, message_type, content)
            VALUES (%s, %s, %s, %s)
            RETURNING id, created_at
        """, (session_id, doc_id, 'user', chat_data.message))
        
        user_msg_result = cursor.fetchone()
        user_timestamp = user_msg_result[1].isoformat()
        
        # Generate AI response
        ai_response = generate_ai_response(chat_data.message, document_text)
        
        # Add AI response to database
        cursor.execute("""
            INSERT INTO chat_messages (session_id, document_id, message_type, content)
            VALUES (%s, %s, %s, %s)
            RETURNING id, created_at
        """, (session_id, doc_id, 'assistant', ai_response))
        
        ai_msg_result = cursor.fetchone()
        ai_timestamp = ai_msg_result[1].isoformat()
        
        # Get recent chat history
        cursor.execute("""
            SELECT message_type, content, created_at
            FROM chat_messages
            WHERE session_id = %s
            ORDER BY created_at DESC
            LIMIT 20
        """, (session_id,))
        
        chat_history = []
        for row in reversed(cursor.fetchall()):
            chat_history.append({
                "role": row[0],
                "content": row[1],
                "timestamp": row[2].isoformat()
            })
        
        conn.commit()
        cursor.close()
        conn.close()
        
        return {
            "response": ai_response,
            "chat_history": chat_history
        }
        
    except HTTPException:
        raise
    except Exception as e:
        print(f"Error in chat: {e}")
        raise HTTPException(status_code=500, detail=f"Chat failed: {str(e)}")

def generate_ai_response(user_message: str, document_text: str) -> str:
    """Generate AI response based on user message and document content using Mistral API"""
    
    if not document_text or document_text.strip() == "":
        return "I don't have access to the document content to answer your question. Please make sure the document has been processed successfully."
    
    # Use Mistral API for intelligent document analysis
    try:
        import requests
        
        api_key = "eyFSYGAUfsrrDmDVLGaKac5IQmFy1gEH"
        base_url = "https://api.mistral.ai/v1/chat/completions"
        
        headers = {
            "Content-Type": "application/json",
            "Authorization": f"Bearer {api_key}"
        }
        
        # Create a focused prompt for document analysis
        system_prompt = """You are an AI assistant specialized in document analysis. You have access to the full text of a document and should provide accurate, specific answers based on the actual content. Always:

1. Answer directly based on the document content
2. Quote specific text when relevant
3. If information isn't in the document, clearly state that
4. Be concise but comprehensive
5. Focus on the user's specific question"""

        user_prompt = f"""Document Content:
{document_text[:8000]}  # Limit to avoid token limits

User Question: {user_message}

Please analyze the document and provide a specific answer to the user's question based on the actual content above."""

        payload = {
            "model": "mistral-large-latest",
            "messages": [
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ],
            "max_tokens": 1000,
            "temperature": 0.3
        }
        
        response = requests.post(base_url, headers=headers, json=payload, timeout=30)
        
        if response.status_code == 200:
            result = response.json()
            return result['choices'][0]['message']['content']
        else:
            print(f"Mistral API error: {response.status_code} - {response.text}")
            return f"I can analyze this document for you. Based on your question about '{user_message}', let me examine the content and provide you with specific information from the document."
            
    except Exception as e:
        print(f"Error calling Mistral API: {e}")
        return f"I can help you analyze this document. Your question about '{user_message}' requires me to examine the document content. Please let me know if you'd like me to focus on any specific sections."

def generate_ai_response_all_docs(user_message: str, all_documents_text: str) -> str:
    """Generate AI response based on user message and all documents content using Mistral API"""
    
    if not all_documents_text or all_documents_text.strip() == "":
        return "I don't have access to any document content to answer your question. Please make sure documents have been processed successfully."
    
    # Use Mistral API for intelligent multi-document analysis
    try:
        import requests
        
        api_key = "eyFSYGAUfsrrDmDVLGaKac5IQmFy1gEH"
        base_url = "https://api.mistral.ai/v1/chat/completions"
        
        headers = {
            "Content-Type": "application/json",
            "Authorization": f"Bearer {api_key}"
        }
        
        # Create a focused prompt for multi-document analysis
        system_prompt = """You are an AI assistant specialized in analyzing multiple documents. You have access to the content of all documents in the library and should provide accurate, specific answers based on the actual content. Always:

1. Answer directly based on the document content
2. When listing documents, provide their actual names from the content
3. Quote specific text when relevant
4. If information isn't in the documents, clearly state that
5. Be concise but comprehensive
6. Focus on the user's specific question
7. When asked about document names or counts, provide the exact information from the content"""

        # Truncate content to avoid token limits while preserving document structure
        truncated_content = all_documents_text[:12000]
        
        user_prompt = f"""All Documents Content:
{truncated_content}

User Question: {user_message}

Please analyze all the documents and provide a specific answer to the user's question based on the actual content above. If the user asks for document names, list them exactly as they appear in the content."""

        payload = {
            "model": "mistral-large-latest",
            "messages": [
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ],
            "max_tokens": 1500,
            "temperature": 0.3
        }
        
        response = requests.post(base_url, headers=headers, json=payload, timeout=30)
        
        if response.status_code == 200:
            result = response.json()
            return result['choices'][0]['message']['content']
        else:
            print(f"Mistral API error: {response.status_code} - {response.text}")
            # Fallback: Parse document names from content
            lines = all_documents_text.split('\n')
            doc_names = []
            for line in lines:
                if 'Document:' in line or 'Filename:' in line:
                    doc_names.append(line.strip())
            
            if doc_names and ("document names" in user_message.lower() or "list" in user_message.lower()):
                return f"Here are the document names I found:\n" + "\n".join(doc_names[:20])  # Limit to first 20
            else:
                return f"I can analyze your document library for you. Based on your question about '{user_message}', let me examine the content and provide you with specific information from the documents."
            
    except Exception as e:
        print(f"Error calling Mistral API: {e}")
        # Fallback: Try to extract document information
        lines = all_documents_text.split('\n')
        doc_count = all_documents_text.count('Document:') or all_documents_text.count('Filename:')
        
        if "how many" in user_message.lower() or "count" in user_message.lower():
            return f"I found {doc_count} documents in your library. I can help you analyze their content or search for specific information."
        elif "document names" in user_message.lower() or "list" in user_message.lower():
            doc_names = []
            for line in lines:
                if 'Document:' in line or 'Filename:' in line:
                    doc_names.append(line.strip())
            if doc_names:
                return f"Here are the document names I found:\n" + "\n".join(doc_names[:20])
            else:
                return f"I have access to {doc_count} documents but need to process them to extract the names. Please try again in a moment."
        else:
            return f"I can help you analyze your document library. Your question about '{user_message}' requires me to examine the document content. Please let me know if you'd like me to focus on any specific aspects."

@app.get("/documents/{doc_id}/chat")
async def get_chat_history(doc_id: str):
    """Get chat history for a document from database"""
    try:
        if not ocr_processor:
            raise HTTPException(status_code=500, detail="OCR processor not available")
            
        conn = ocr_processor.get_db_connection()
        cursor = conn.cursor()
        
        # Check if document exists
        cursor.execute("SELECT id FROM documents WHERE id = %s", (doc_id,))
        if not cursor.fetchone():
            cursor.close()
            conn.close()
            raise HTTPException(status_code=404, detail="Document not found")
        
        # Get chat history
        cursor.execute("""
            SELECT message_type, content, created_at
            FROM chat_messages
            WHERE document_id = %s
            ORDER BY created_at ASC
        """, (doc_id,))
        
        chat_history = []
        for row in cursor.fetchall():
            chat_history.append({
                "role": row[0],
                "content": row[1],
                "timestamp": row[2].isoformat()
            })
        
        cursor.close()
        conn.close()
        
        return chat_history
        
    except HTTPException:
        raise
    except Exception as e:
        print(f"Error getting chat history: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to get chat history: {str(e)}")

@app.delete("/documents/{doc_id}")
async def delete_document(doc_id: str):
    """Delete a document from database"""
    try:
        if not ocr_processor:
            raise HTTPException(status_code=500, detail="OCR processor not available")
            
        conn = ocr_processor.get_db_connection()
        cursor = conn.cursor()
        
        # Get document file path before deletion
        cursor.execute("SELECT file_path FROM documents WHERE id = %s", (doc_id,))
        result = cursor.fetchone()
        
        if not result:
            cursor.close()
            conn.close()
            raise HTTPException(status_code=404, detail="Document not found")
        
        file_path = Path(result[0]) if result[0] else None
        
        # Delete from database first (CASCADE will handle related records)
        cursor.execute("DELETE FROM documents WHERE id = %s", (doc_id,))
        conn.commit()
        cursor.close()
        conn.close()
        
        # Try to remove physical files (non-critical if fails)
        file_deletion_errors = []
        
        if file_path and file_path.exists():
            try:
                file_path.unlink()
            except (PermissionError, OSError) as e:
                file_deletion_errors.append(f"Could not delete file {file_path.name}: {str(e)}")
                print(f"Warning: Could not delete file {file_path}: {e}")
        
        # Remove output file if exists
        output_path = Path(f"outputs/{doc_id}_processed.docx")
        if output_path.exists():
            try:
                output_path.unlink()
            except (PermissionError, OSError) as e:
                file_deletion_errors.append(f"Could not delete output file: {str(e)}")
                print(f"Warning: Could not delete output file {output_path}: {e}")
        
        # Return success message with warnings if any
        message = "Document deleted successfully"
        if file_deletion_errors:
            message += f" (Note: Some files could not be deleted: {'; '.join(file_deletion_errors)})"
        
        return {"message": message}
        
    except HTTPException:
        raise
    except Exception as e:
        print(f"Error deleting document: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to delete document: {str(e)}")

@app.get("/documents/{doc_id}/export")
async def export_document(doc_id: str, format: str = "txt"):
    """Export document in specified format"""
    try:
        if not ocr_processor:
            raise HTTPException(status_code=500, detail="OCR processor not available")
            
        # Get document info from database
        conn = ocr_processor.get_db_connection()
        cursor = conn.cursor()
        
        cursor.execute("""
            SELECT original_filename, filename
            FROM documents 
            WHERE id = %s
        """, (doc_id,))
        
        doc_result = cursor.fetchone()
        if not doc_result:
            cursor.close()
            conn.close()
            raise HTTPException(status_code=404, detail="Document not found")
        
        original_filename, filename = doc_result
        base_name = original_filename.replace('.pdf', '') if original_filename else filename.replace('.pdf', '')
        
        if format == "txt":
            # Prefer validated text; fallback to complete_document
            cursor.execute("""
                SELECT processed_text 
                FROM extracted_content 
                WHERE document_id = %s AND content_type = 'validated_document'
                ORDER BY updated_at DESC NULLS LAST, created_at DESC
                LIMIT 1
            """, (doc_id,))
            row = cursor.fetchone()
            complete_text = row[0] if row and row[0] else ""
            
            if not complete_text:
                cursor.execute("""
                    SELECT raw_text 
                    FROM extracted_content 
                    WHERE document_id = %s AND content_type = 'complete_document'
                    ORDER BY created_at DESC
                    LIMIT 1
                """, (doc_id,))
                result = cursor.fetchone()
                complete_text = result[0] if result else ""
            
            # If no complete document found, try to get all page content
            if not complete_text:
                cursor.execute("""
                    SELECT raw_text 
                    FROM extracted_content 
                    WHERE document_id = %s AND page_id IS NOT NULL
                    ORDER BY metadata->>'page_number'
                """, (doc_id,))
                
                page_results = cursor.fetchall()
                if page_results:
                    complete_text = "\n\n".join([row[0] for row in page_results if row[0]])
            
            cursor.close()
            conn.close()
            
            if not complete_text:
                raise HTTPException(status_code=404, detail="No extracted text found for this document")
            
            # Create text file
            export_path = OUTPUT_DIR / f"{doc_id}_export.txt"
            with open(export_path, "w", encoding="utf-8") as f:
                f.write(complete_text)
            
            return FileResponse(
                path=export_path,
                filename=f"{base_name}.txt",
                media_type="text/plain"
            )
        
        elif format == "docx":
            cursor.close()
            conn.close()
            
            # Look for existing DOCX file in outputs directory
            # Try different possible naming patterns
            possible_docx_files = [
                OUTPUT_DIR / f"{doc_id}_{base_name}_extracted.docx",
                OUTPUT_DIR / f"{doc_id}_{filename.replace('.pdf', '')}_extracted.docx",
                OUTPUT_DIR / f"{filename.replace('.pdf', '')}_extracted.docx"
            ]
            
            # Find existing DOCX file
            docx_file = None
            for possible_file in possible_docx_files:
                if possible_file.exists():
                    docx_file = possible_file
                    break
            
            # If no existing DOCX found, look for any DOCX file with the doc_id
            if not docx_file:
                for file_path in OUTPUT_DIR.glob(f"*{doc_id}*.docx"):
                    docx_file = file_path
                    break
            
            if docx_file and docx_file.exists():
                return FileResponse(
                    path=docx_file,
                    filename=f"{base_name}.docx",
                    media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            else:
                # Generate DOCX on the fly from DB extracted text
                try:
                    from docx import Document as DocxDocument
                except Exception:
                    # Lazy install if missing in the environment
                    try:
                        import subprocess, sys as _sys
                        subprocess.check_call([_sys.executable, "-m", "pip", "install", "python-docx"], stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
                        from docx import Document as DocxDocument
                    except Exception as _e:
                        raise HTTPException(status_code=500, detail=f"DOCX generation unavailable: {_e}")

                # Re-open DB to fetch text
                conn2 = ocr_processor.get_db_connection()
                cur2 = conn2.cursor()
                cur2.execute("""
                    SELECT raw_text 
                    FROM extracted_content 
                    WHERE document_id = %s AND content_type = 'complete_document'
                    ORDER BY created_at DESC
                    LIMIT 1
                """, (doc_id,))
                row = cur2.fetchone()
                complete_text = row[0] if row else ""

                if not complete_text:
                    # Fallback: stitch page texts
                    cur2.execute("""
                        SELECT raw_text 
                        FROM extracted_content 
                        WHERE document_id = %s AND page_id IS NOT NULL
                        ORDER BY metadata->>'page_number'
                    """, (doc_id,))
                    page_rows = cur2.fetchall()
                    if page_rows:
                        complete_text = "\n\n".join([r[0] for r in page_rows if r and r[0]])

                cur2.close()
                conn2.close()

                if not complete_text:
                    raise HTTPException(status_code=404, detail="No extracted text available to generate DOCX")

                # Build DOCX
                doc = DocxDocument()
                for paragraph in complete_text.split("\n\n"):
                    doc.add_paragraph(paragraph)

                OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
                generated_path = OUTPUT_DIR / f"{doc_id}_{base_name}_extracted.docx"
                doc.save(str(generated_path))

                return FileResponse(
                    path=generated_path,
                    filename=f"{base_name}.docx",
                    media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
        
        else:
            raise HTTPException(status_code=400, detail="Unsupported export format. Use 'txt' or 'docx'.")
            
    except HTTPException:
        raise
    except Exception as e:
        print(f"Error exporting document: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to export document: {str(e)}")

@app.get("/documents/{doc_id}/pages/{page_number}/image")
async def get_page_image(doc_id: str, page_number: int):
    """Serve a specific page image as binary data"""
    try:
        if not ocr_processor:
            raise HTTPException(status_code=500, detail="OCR processor not available")
            
        conn = ocr_processor.get_db_connection()
        cursor = conn.cursor()
        
        # Get the specific page image
        cursor.execute("""
            SELECT image_data
            FROM pages 
            WHERE document_id = %s AND page_number = %s
        """, (doc_id, page_number))
        
        result = cursor.fetchone()
        cursor.close()
        conn.close()
        
        if not result or not result[0]:
            raise HTTPException(status_code=404, detail="Page image not found")
        
        image_data = result[0]
        
        # Convert memoryview to bytes if necessary
        if isinstance(image_data, memoryview):
            image_data = image_data.tobytes()
        
        # Return the image as binary data
        from fastapi.responses import Response
        return Response(
            content=image_data,
            media_type="image/jpeg",
            headers={
                "Cache-Control": "public, max-age=3600",  # Cache for 1 hour
                "Content-Disposition": f"inline; filename=page_{page_number}.jpg"
            }
        )
        
    except Exception as e:
        print(f"Error serving page image: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to serve page image: {str(e)}")

@app.get("/documents/{doc_id}/pages")
async def get_document_pages(doc_id: str):
    """Get document pages with images for validation"""
    try:
        if not ocr_processor:
            raise HTTPException(status_code=500, detail="OCR processor not available")
            
        conn = ocr_processor.get_db_connection()
        cursor = conn.cursor()
        
        # Get document pages with extracted content
        cursor.execute("""
            SELECT p.id, p.page_number, p.image_data, 
                   COALESCE(ec.raw_text, ec.processed_text, '') as extracted_text
            FROM pages p
            LEFT JOIN extracted_content ec ON p.id = ec.page_id
            WHERE p.document_id = %s
            ORDER BY p.page_number
        """, (doc_id,))
        
        pages = []
        fetched_rows = cursor.fetchall()
        for row in fetched_rows:
            page_id, page_number, image_data, extracted_text = row
            
            # Convert image data to base64 for frontend display
            image_url = None
            if image_data:
                import base64
                # Handle both raw bytes and already-base64-encoded strings stored previously
                try:
                    # If it's a memoryview, convert to bytes first
                    if isinstance(image_data, memoryview):
                        image_bytes = image_data.tobytes()
                    else:
                        image_bytes = image_data

                    # Heuristic: if the bytes look like ASCII base64 (only base64 charset and '=' padding)
                    is_ascii = False
                    try:
                        sample = image_bytes[:64]
                        sample.decode('ascii')
                        is_ascii = True
                    except Exception:
                        is_ascii = False

                    if is_ascii:
                        text = image_bytes.decode('ascii', errors='ignore')
                        if len(text.strip()) > 0 and all(c in 'ABCDEFGHIJKLMNOPQRSTUVWXYZabcdefghijklmnopqrstuvwxyz0123456789+/=\n\r' for c in text.strip()):
                            # Assume it's already base64 text stored earlier; use as-is
                            image_base64 = text.replace('\n', '').replace('\r', '')
                        else:
                            image_base64 = base64.b64encode(image_bytes).decode('utf-8')
                    else:
                        image_base64 = base64.b64encode(image_bytes).decode('utf-8')

                    image_url = f"data:image/jpeg;base64,{image_base64}"
                except Exception as img_err:
                    print(f"[WARN] Failed to prepare image data for page {page_number}: {img_err}")
                    image_url = None
            
            pages.append({
                "id": str(page_id),
                "pageNumber": page_number,
                "imageUrl": image_url,
                "extractedText": extracted_text or "",
                "validated": False  # TODO: Add validation status to database
            })

        # If no pages found, dynamically render from original PDF as a resilience fallback
        if len(pages) == 0:
            try:
                # Get original file_path
                cursor.execute("SELECT file_path, total_pages FROM documents WHERE id = %s", (doc_id,))
                doc_row = cursor.fetchone()
                file_path = doc_row[0] if doc_row else None
                if not file_path or not Path(file_path).exists():
                    cursor.close()
                    conn.close()
                    return pages

                # Try PyMuPDF first
                rendered_any = False
                try:
                    import fitz  # PyMuPDF
                    pdf_doc = fitz.open(file_path)
                    for page_num in range(len(pdf_doc)):
                        page = pdf_doc.load_page(page_num)
                        mat = fitz.Matrix(2, 2)
                        pix = None
                        if hasattr(page, 'get_pixmap'):
                            pix = getattr(page, 'get_pixmap')(matrix=mat)
                        elif hasattr(page, 'getPixmap'):
                            pix = getattr(page, 'getPixmap')(matrix=mat)
                        else:
                            pix = getattr(page, 'get_pixmap')() if hasattr(page, 'get_pixmap') else getattr(page, 'getPixmap')()
                        img_bytes = pix.tobytes("jpeg")
                        # Save for future requests
                        cursor.execute(
                            """
                            INSERT INTO pages (id, document_id, page_number, page_type, image_data, processing_status, created_at)
                            VALUES (%s, %s, %s, %s, %s, %s, %s)
                            ON CONFLICT (document_id, page_number) DO UPDATE SET image_data = EXCLUDED.image_data
                            """,
                            (str(uuid.uuid4()), doc_id, page_num + 1, 'standard', img_bytes, 'completed', datetime.now())
                        )
                        import base64
                        image_base64 = base64.b64encode(img_bytes).decode('utf-8')
                        pages.append({
                            "id": None,
                            "pageNumber": page_num + 1,
                            "imageUrl": f"data:image/jpeg;base64,{image_base64}",
                            "extractedText": "",
                            "validated": False
                        })
                        rendered_any = True
                    pdf_doc.close()
                except Exception:
                    rendered_any = False

                # Fallback to pdf2image if PyMuPDF path failed
                if not rendered_any:
                    try:
                        from pdf2image import convert_from_path
                        import io, base64
                        images = convert_from_path(file_path, dpi=150, fmt='JPEG')
                        for idx, page_image in enumerate(images, 1):
                            buf = io.BytesIO()
                            page_image.save(buf, format='JPEG', quality=85)
                            img_bytes = buf.getvalue()
                            cursor.execute(
                                """
                                INSERT INTO pages (id, document_id, page_number, page_type, image_data, processing_status, created_at)
                                VALUES (%s, %s, %s, %s, %s, %s, %s)
                                ON CONFLICT (document_id, page_number) DO UPDATE SET image_data = EXCLUDED.image_data
                                """,
                                (str(uuid.uuid4()), doc_id, idx, 'standard', img_bytes, 'completed', datetime.now())
                            )
                            image_base64 = base64.b64encode(img_bytes).decode('utf-8')
                            pages.append({
                                "id": None,
                                "pageNumber": idx,
                                "imageUrl": f"data:image/jpeg;base64,{image_base64}",
                                "extractedText": "",
                                "validated": False
                            })
                    except Exception:
                        # If all rendering fails, return empty pages gracefully
                        pass

                conn.commit()
            except Exception as gen_err:
                print(f"[WARN] Dynamic page generation failed for {doc_id}: {gen_err}")

        cursor.close()
        conn.close()
        
        return pages
        
    except Exception as e:
        print(f"Error getting document pages: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to get document pages: {str(e)}")

@app.post("/documents/{doc_id}/pages/{page_id}/validate")
async def validate_page(doc_id: str, page_id: str, data: ValidateText):
    """Validate and save text for a specific page"""
    try:
        if not ocr_processor:
            raise HTTPException(status_code=500, detail="OCR processor not available")
            
        conn = ocr_processor.get_db_connection()
        cursor = conn.cursor()
        
        # Update the extracted content with validated text
        cursor.execute("""
            UPDATE extracted_content 
            SET processed_text = %s, updated_at = CURRENT_TIMESTAMP
            WHERE page_id = %s AND document_id = %s
        """, (data.validated_text, page_id, doc_id))
        
        if cursor.rowcount == 0:
            # If no existing content, insert new record
            cursor.execute("""
                INSERT INTO extracted_content (page_id, document_id, raw_text, processed_text)
                VALUES (%s, %s, %s, %s)
            """, (page_id, doc_id, data.validated_text, data.validated_text))
        
        conn.commit()
        cursor.close()
        conn.close()
        
        return {"message": "Page validated successfully"}
        
    except Exception as e:
        print(f"Error validating page: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to validate page: {str(e)}")

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000, reload=True)