import os
import requests
import json
import base64
import uuid
from datetime import datetime
from io import BytesIO
from pathlib import Path
from pdf2image import convert_from_path
from PIL import Image
from typing import List, Dict
import time
import subprocess
import sys
import psycopg2
from psycopg2.extras import RealDictCursor
from fastapi import FastAPI, File, UploadFile, HTTPException
from fastapi.middleware.cors import CORSMiddleware
import shutil

class DatabaseOCR:
    def __init__(self, api_key: str = "eyFSYGAUfsrrDmDVLGaKac5IQmFy1gEH"):
        """
        Initialize Database OCR client with Mistral Pixtral API integration
        """
        # Mistral API Configuration
        self.api_key = api_key
        self.base_url = "https://api.mistral.ai/v1/chat/completions"
        self.model = "pixtral-12b-2409"
        
        # Headers for API requests
        self.headers = {
            "Content-Type": "application/json",
            "Authorization": f"Bearer {self.api_key}"
        }
        
        # Database Configuration
        self.db_config = {
            'host': '127.0.0.1',
            'port': 5432,
            'database': 'LLMAPI',
            'user': 'postgres',
            'password': 'shibin'
        }
        
        # Ensure required packages are installed
        self.ensure_packages_installed()
        
        # Initialize database
        self.init_database()

    def ensure_packages_installed(self):
        """Ensure required packages are installed"""
        packages = ['python-docx', 'psycopg2-binary', 'fastapi', 'python-multipart']
        
        for package in packages:
            try:
                if package == 'python-docx':
                    import docx
                elif package == 'psycopg2-binary':
                    import psycopg2
                elif package == 'fastapi':
                    import fastapi
                elif package == 'python-multipart':
                    import multipart
            except ImportError:
                print(f"Installing {package}...")
                subprocess.check_call([sys.executable, "-m", "pip", "install", package])
                print(f"✅ {package} installed successfully!")

    def get_db_connection(self):
        """Get database connection"""
        try:
            conn = psycopg2.connect(**self.db_config)
            return conn
        except Exception as e:
            print(f"❌ Database connection error: {e}")
            raise
    
    def init_database(self):
        """Initialize database tables for OCR processing"""
        conn = self.get_db_connection()
        cursor = conn.cursor()
        
        try:
            # Create documents table if not exists
            cursor.execute("""
                CREATE TABLE IF NOT EXISTS documents (
                    id UUID PRIMARY KEY DEFAULT gen_random_uuid(),
                    filename VARCHAR(255) NOT NULL,
                    original_filename VARCHAR(255) NOT NULL,
                    file_size BIGINT,
                    file_path TEXT,
                    mime_type VARCHAR(100),
                    upload_date TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                    processing_status VARCHAR(50) DEFAULT 'pending',
                    total_pages INTEGER,
                    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                    updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                )
            """)
            
            # Create pages table if not exists
            cursor.execute("""
                CREATE TABLE IF NOT EXISTS pages (
                    id UUID PRIMARY KEY DEFAULT gen_random_uuid(),
                    document_id UUID REFERENCES documents(id) ON DELETE CASCADE,
                    page_number INTEGER NOT NULL,
                    page_type VARCHAR(50) DEFAULT 'text',
                    image_data BYTEA,
                    processing_status VARCHAR(50) DEFAULT 'pending',
                    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                    updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                )
            """)
            
            # Create extracted_content table if not exists
            cursor.execute("""
                CREATE TABLE IF NOT EXISTS extracted_content (
                    id UUID PRIMARY KEY DEFAULT gen_random_uuid(),
                    page_id UUID REFERENCES pages(id) ON DELETE CASCADE,
                    document_id UUID REFERENCES documents(id) ON DELETE CASCADE,
                    content_type VARCHAR(50) DEFAULT 'text',
                    raw_text TEXT,
                    processed_text TEXT,
                    confidence_score FLOAT,
                    metadata JSONB,
                    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                    updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                )
            """)
            
            # Create processing_logs table if not exists
            cursor.execute("""
                CREATE TABLE IF NOT EXISTS processing_logs (
                    id UUID PRIMARY KEY DEFAULT gen_random_uuid(),
                    document_id UUID REFERENCES documents(id) ON DELETE CASCADE,
                    page_id UUID REFERENCES pages(id) ON DELETE CASCADE,
                    step_name VARCHAR(100),
                    status VARCHAR(50),
                    message TEXT,
                    processing_time_ms INTEGER,
                    error_details JSONB,
                    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                )
            """)
            
            conn.commit()
            print("✅ Database tables initialized successfully")
            
        except Exception as e:
            conn.rollback()
            print(f"❌ Database initialization error: {e}")
            raise
        finally:
            cursor.close()
            conn.close()
    
    def test_database_connection(self) -> Dict:
        """Test database connection"""
        try:
            conn = self.get_db_connection()
            cursor = conn.cursor()
            cursor.execute("SELECT 1")
            result = cursor.fetchone()
            cursor.close()
            conn.close()
            
            if result:
                return {"success": True, "message": "Database connection successful"}
            else:
                return {"success": False, "error": "Database query failed"}
                
        except Exception as e:
            return {"success": False, "error": f"Database connection failed: {str(e)}"}

    def pdf_to_images(self, pdf_path: str, dpi: int = 300) -> List[Image.Image]:
        """Convert PDF to images"""
        print(f"Converting PDF to images: {pdf_path}")
        try:
            images = convert_from_path(pdf_path, dpi=dpi)
            print(f"Successfully converted {len(images)} pages")
            return images
        except Exception as e:
            # Enhanced error message for poppler
            print(f"Error converting PDF: {e}. Make sure 'poppler-utils' is installed.")
            return []


    def image_to_base64(self, image: Image.Image) -> str:
        """Convert PIL Image to base64 string"""
        buffer = BytesIO()
        image.save(buffer, format='PNG')
        img_bytes = buffer.getvalue()
        return base64.b64encode(img_bytes).decode('utf-8')

    def save_document_to_db(self, filename: str, original_filename: str, file_size: int, file_path: str, mime_type: str) -> str:
        """Save document metadata to database"""
        conn = self.get_db_connection()
        cursor = conn.cursor(cursor_factory=RealDictCursor)
        
        try:
            cursor.execute("""
                INSERT INTO documents (filename, original_filename, file_size, file_path, mime_type)
                VALUES (%s, %s, %s, %s, %s)
                RETURNING id
            """, (filename, original_filename, file_size, file_path, mime_type))
            
            result = cursor.fetchone()
            document_id = result['id']
            conn.commit()
            return str(document_id)
            
        except Exception as e:
            conn.rollback()
            print(f"❌ Error saving document to database: {e}")
            raise
        finally:
            cursor.close()
            conn.close()
    
    def save_page_to_db(self, document_id: str, page_number: int, image_data: bytes) -> str:
        """Save page data to database"""
        conn = self.get_db_connection()
        cursor = conn.cursor(cursor_factory=RealDictCursor)
        
        try:
            cursor.execute("""
                INSERT INTO pages (document_id, page_number, image_data)
                VALUES (%s, %s, %s)
                RETURNING id
            """, (document_id, page_number, image_data))
            
            result = cursor.fetchone()
            page_id = result['id']
            conn.commit()
            return str(page_id)
            
        except Exception as e:
            conn.rollback()
            print(f"❌ Error saving page to database: {e}")
            raise
        finally:
            cursor.close()
            conn.close()
    
    def save_extracted_content_to_db(self, page_id: str, document_id: str, text: str, metadata: dict = None) -> str:
        """Save extracted content to database"""
        conn = self.get_db_connection()
        cursor = conn.cursor(cursor_factory=RealDictCursor)
        
        try:
            cursor.execute("""
                INSERT INTO extracted_content 
                (page_id, document_id, raw_text, processed_text, metadata)
                VALUES (%s, %s, %s, %s, %s)
                RETURNING id
            """, (page_id, document_id, text, text, 
                  json.dumps(metadata) if metadata else None))
            
            result = cursor.fetchone()
            content_id = result['id']
            conn.commit()
            return str(content_id)
            
        except Exception as e:
            conn.rollback()
            print(f"❌ Error saving extracted content to database: {e}")
            raise
        finally:
            cursor.close()
            conn.close()
    
    def log_processing_step(self, document_id: str, page_id: str, step_name: str, status: str, message: str = None, processing_time_ms: int = None, error_details: dict = None):
        """Log processing step to database"""
        conn = self.get_db_connection()
        cursor = conn.cursor()
        
        try:
            cursor.execute("""
                INSERT INTO processing_logs 
                (document_id, page_id, step_name, status, message, processing_time_ms, error_details)
                VALUES (%s, %s, %s, %s, %s, %s, %s)
            """, (document_id, page_id, step_name, status, message, processing_time_ms, 
                  json.dumps(error_details) if error_details else None))
            
            conn.commit()
            
        except Exception as e:
            conn.rollback()
            print(f"❌ Error logging processing step: {e}")
        finally:
            cursor.close()
            conn.close()
    
    def update_document_status(self, document_id: str, status: str, total_pages: int = None):
        """Update document processing status"""
        conn = self.get_db_connection()
        cursor = conn.cursor()
        
        try:
            if total_pages is not None:
                cursor.execute("""
                    UPDATE documents 
                    SET processing_status = %s, total_pages = %s, updated_at = CURRENT_TIMESTAMP
                    WHERE id = %s
                """, (status, total_pages, document_id))
            else:
                cursor.execute("""
                    UPDATE documents 
                    SET processing_status = %s, updated_at = CURRENT_TIMESTAMP
                    WHERE id = %s
                """, (status, document_id))
            
            conn.commit()
            
        except Exception as e:
            conn.rollback()
            print(f"❌ Error updating document status: {e}")
            raise
        finally:
            cursor.close()
            conn.close()
    
    def test_api_connection(self) -> Dict:
        """
        Test Mistral API connection with a simple request
        """
        test_payload = {
            "model": self.model,
            "messages": [{"role": "user", "content": "Hello, test connection"}],
            "max_tokens": 10
        }

        try:
            response = requests.post(
                self.base_url,
                headers=self.headers,
                json=test_payload,
                timeout=30
            )

            if response.status_code == 200:
                return {"success": True, "message": "API connection successful"}
            else:
                return {"success": False, "error": f"API Error {response.status_code}: {response.text}"}

        except Exception as e:
            return {"success": False, "error": f"Connection failed: {str(e)}"}

    def transcribe_image(self, image: Image.Image, custom_prompt: str = None) -> Dict:
        """Transcribe single image using Mistral Pixtral API"""
        base64_image = self.image_to_base64(image)
        prompt = custom_prompt or """
        Transcribe ALL text from this image with maximum accuracy. Follow these guidelines:
        1. Extract every visible text element including headers, body text, footnotes
        2. Preserve original formatting and layout structure
        3. Maintain paragraph breaks and line spacing
        4. Include tables, lists, and any formatted content
        5. Transcribe numbers, dates, and special characters accurately
        6. If text is handwritten, read carefully and indicate if unclear
        7. Process multi-column layouts from left to right

        Provide the transcribed text maintaining the document structure.
        """
        
        payload = {
            "model": self.model,
            "messages": [
                {
                    "role": "user",
                    "content": [
                        {"type": "text", "text": prompt},
                        {"type": "image_url", "image_url": f"data:image/png;base64,{base64_image}"}
                    ]
                }
            ],
            "max_tokens": 4000,
            "temperature": 0.1
        }
        
        try:
            response = requests.post(
                self.base_url,
                headers=self.headers,
                json=payload,
                timeout=60
            )
            
            if response.status_code == 200:
                result = response.json()
                return {
                    "success": True,
                    "text": result['choices'][0]['message']['content'],
                    "usage": result.get('usage', {}),
                    "model": result.get('model', self.model)
                }
            else:
                error_detail = ""
                try:
                    error_data = response.json()
                    error_detail = error_data.get('message', response.text)
                except:
                    error_detail = response.text
                return {"success": False, "error": f"API Error {response.status_code}: {error_detail}"}
                
        except Exception as e:
            return {"success": False, "error": f"Request failed: {str(e)}"}

    def create_word_document(self, results: List[Dict], output_path: Path, filename: str) -> Path:
        """Create a Word document from extraction results with enhanced formatting"""
        try:
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.oxml.shared import OxmlElement, qn

            doc = Document()

            # Set document margins
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
                section.left_margin = Inches(1)
                section.right_margin = Inches(1)

            # Add title with formatting
            title = doc.add_heading(f'OCR Extraction Results', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Add subtitle
            subtitle = doc.add_heading(f'Document: {filename}', level=1)
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Add metadata section
            doc.add_paragraph()  # Empty line
            meta_heading = doc.add_heading('Document Information', level=2)

            # Create metadata table
            meta_table = doc.add_table(rows=4, cols=2)
            meta_table.style = 'Table Grid'

            # Fill metadata
            meta_table.cell(0, 0).text = 'Extraction Method'
            meta_table.cell(0, 1).text = 'Mistral Pixtral API'
            meta_table.cell(1, 0).text = 'Total Pages'
            meta_table.cell(1, 1).text = str(sum(1 for r in results if r["status"] == "success" or r["status"] == "error")) # Use count from results to be accurate
            meta_table.cell(2, 0).text = 'Successfully Processed'
            meta_table.cell(2, 1).text = str(sum(1 for r in results if r["status"] == "success"))
            meta_table.cell(3, 0).text = 'Processing Date'
            meta_table.cell(3, 1).text = time.strftime("%Y-%m-%d %H:%M:%S")

            doc.add_paragraph()  # Empty line
            doc.add_paragraph("-" * 80)

            # Add extracted content
            for result in results:
                if result["status"] == "success" and result["text"]:
                    # Add page header
                    page_heading = doc.add_heading(f'Page {result["page"]}', level=2)

                    # Add character count info
                    info_para = doc.add_paragraph(f'Characters extracted: {result.get("text_length", len(result["text"]))}')
                    info_para.style = 'Intense Quote'

                    # Process and add text content
                    text_content = result["text"].strip()

                    # Split into paragraphs and add them
                    paragraphs = text_content.split('\n')
                    current_paragraph = ""

                    for line in paragraphs:
                        line = line.strip()
                        if line:
                            if current_paragraph:
                                current_paragraph += " " + line
                            else:
                                current_paragraph = line
                        else:
                            # Empty line indicates paragraph break
                            if current_paragraph:
                                para = doc.add_paragraph(current_paragraph)
                                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                current_paragraph = ""

                    # Add any remaining text
                    if current_paragraph:
                        para = doc.add_paragraph(current_paragraph)
                        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

                    # Add page separator except for last page
                    if result["page"] < len(results):
                        doc.add_page_break()

                elif result["status"] == "error":
                    # Add error information
                    error_heading = doc.add_heading(f'Page {result["page"]} - Processing Error', level=2)
                    error_para = doc.add_paragraph(f'Error: {result.get("error", "Unknown error")}')
                    error_para.style = 'Intense Quote'

            # Save document
            word_file = output_path / f"{filename}_extracted.docx"
            doc.save(str(word_file))
            print(f"✅ Word document created successfully: {word_file}")
            return word_file

        except ImportError as e:
            print(f"❌ python-docx not available: {e}")
            print("Creating fallback text file...")
            return self.create_fallback_document(results, output_path, filename)
        except Exception as e:
            print(f"❌ Error creating Word document: {e}")
            print("Creating fallback text file...")
            return self.create_fallback_document(results, output_path, filename)


    def create_fallback_document(self, results: List[Dict], output_path: Path, filename: str) -> Path:
        """Create a formatted text file as fallback"""
        fallback_file = output_path / f"{filename}_extracted_fallback.txt"

        with open(fallback_file, 'w', encoding='utf-8') as f:
            f.write("="*80 + "\n")
            f.write(f"OCR EXTRACTION RESULTS\n")
            f.write(f"Document: {filename}\n")
            f.write("="*80 + "\n\n")

            f.write("DOCUMENT INFORMATION:\n")
            f.write("-" * 40 + "\n")
            f.write(f"Extraction Method: Mistral Pixtral API\n")
            f.write(f"Total Pages: {len(results)}\n")
            f.write(f"Successfully Processed: {sum(1 for r in results if r['status'] == 'success')}\n")
            f.write(f"Processing Date: {time.strftime('%Y-%m-%d %H:%M:%S')}\n\n")

            for result in results:
                if result["status"] == "success" and result["text"]:
                    f.write("="*80 + "\n")
                    f.write(f"PAGE {result['page']}\n")
                    f.write("="*80 + "\n\n")
                    f.write(result["text"])
                    f.write("\n\n")
                elif result["status"] == "error":
                    f.write("="*80 + "\n")
                    f.write(f"PAGE {result['page']} - ERROR\n")
                    f.write("="*80 + "\n")
                    f.write(f"Error: {result.get('error', 'Unknown error')}\n\n")

        print(f"✅ Fallback text document created: {fallback_file}")
        return fallback_file


    def process_pdf_from_frontend(self, pdf_path: str, original_filename: str, output_dir: str = "outputs") -> Dict:
        """Complete PDF text extraction pipeline with database integration"""
        start_time = datetime.now()
        print(f"Starting PDF text extraction: {pdf_path}")

        # Test API and database connections first
        api_test = self.test_api_connection()
        if not api_test["success"]:
            return {"success": False, "error": f"Mistral API connection failed: {api_test['error']}"}
            
        db_test = self.test_database_connection()
        if not db_test["success"]:
            return {"success": False, "error": f"Database connection failed: {db_test['error']}"}

        # Create output directory
        output_path = Path(output_dir)
        output_path.mkdir(exist_ok=True)

        # Save document to database
        file_size = os.path.getsize(pdf_path)
        document_id = self.save_document_to_db(
            filename=os.path.basename(pdf_path),
            original_filename=original_filename,
            file_size=file_size,
            file_path=pdf_path,
            mime_type='application/pdf'
        )
        
        print(f"📄 Document saved to database with ID: {document_id}")
        self.log_processing_step(document_id, None, 'document_upload', 'completed', f"Document {original_filename} uploaded successfully")

        # Convert PDF to images
        images = self.pdf_to_images(pdf_path)
        if not images:
            self.update_document_status(document_id, 'failed')
            self.log_processing_step(document_id, None, 'pdf_conversion', 'failed', "Failed to convert PDF to images")
            return {"success": False, "error": "Failed to convert PDF to images. Check poppler-utils installation."}

        # Update document with total pages
        self.update_document_status(document_id, 'processing', len(images))
        self.log_processing_step(document_id, None, 'pdf_conversion', 'completed', f"Converted {len(images)} pages")

        # Extract text from each page
        extraction_results = []
        total_text = ""
        successful_pages = 0
        total_tokens_used = 0
        all_extracted_content = []

        for i, image in enumerate(images, 1):
            page_start_time = datetime.now()
            print(f"Processing page {i}/{len(images)}...")

            # Save page image to database
            img_bytes = BytesIO()
            image.save(img_bytes, format='PNG')
            img_bytes.seek(0)
            image_data = img_bytes.getvalue()
            img_bytes.close()
            
            page_id = self.save_page_to_db(document_id, i, image_data)
            
            # OCR the page using Mistral Pixtral API
            result = self.transcribe_image(image)

            if result["success"]:
                page_text = result["text"]
                total_text += f"\n\n{'='*50}\nPAGE {i}\n{'='*50}\n\n{page_text}"
                successful_pages += 1

                # Track token usage
                if 'usage' in result:
                    total_tokens_used += result['usage'].get('total_tokens', 0)

                # Save extracted content to database
                content_id = self.save_extracted_content_to_db(
                    page_id=page_id,
                    document_id=document_id,
                    text=page_text,
                    metadata={'page_number': i, 'extraction_method': 'mistral_pixtral_api', 'model': result.get('model', self.model)}
                )

                # Save individual page file
                page_file = output_path / f"page_{i:03d}.txt"
                with open(page_file, 'w', encoding='utf-8') as f:
                    f.write(page_text)

                extraction_results.append({
                    "page": i,
                    "text": page_text,
                    "status": "success",
                    "usage": result.get("usage", {}),
                    "text_length": len(page_text),
                    "page_id": page_id,
                    "content_id": content_id
                })
                
                all_extracted_content.append(f"Page {i}: {page_text}")

                # Log successful page processing
                page_time = (datetime.now() - page_start_time).total_seconds() * 1000
                self.log_processing_step(document_id, page_id, 'page_processing', 'completed', 
                                       f"Page {i} processed successfully", int(page_time))
                
                print(f"✅ Page {i} processed successfully ({len(page_text)} characters)")

            else:
                print(f"❌ Error on page {i}: {result['error']}")
                extraction_results.append({
                    "page": i,
                    "text": "",
                    "status": "error",
                    "error": result["error"],
                    "page_id": page_id
                })
                
                # Log page processing error
                page_time = (datetime.now() - page_start_time).total_seconds() * 1000
                self.log_processing_step(document_id, page_id, 'page_processing', 'failed', 
                                       f"Page {i} processing failed: {result['error']}", int(page_time),
                                       {'error': result['error']})

            # Rate limiting - be respectful to Mistral API
            time.sleep(2)

        # Create Word document
        word_file = self.create_word_document(extraction_results, output_path, Path(pdf_path).stem)

        # Save complete text
        complete_file = output_path / f"{Path(pdf_path).stem}_complete.txt"
        with open(complete_file, 'w', encoding='utf-8') as f:
            f.write(total_text)
        
        # Save complete document text to database
        complete_text = "\n\n".join(all_extracted_content)
        if complete_text:
            self.save_extracted_content_to_db(
                page_id=None,
                document_id=document_id,
                text=complete_text,
                metadata={'content_type': 'complete_document', 'total_pages': len(images), 'extraction_method': 'mistral_pixtral_api'}
            )

        # Update final document status
        final_status = 'completed' if successful_pages > 0 else 'failed'
        self.update_document_status(document_id, final_status)
        
        # Log completion
        total_time = (datetime.now() - start_time).total_seconds() * 1000
        self.log_processing_step(document_id, None, 'complete_processing', 'completed', 
                               f"Document processing completed", int(total_time))

        # Prepare summary
        summary = {
            "success": True,
            "document_id": document_id,
            "pdf_file": pdf_path,
            "total_pages": len(images),
            "successful_pages": successful_pages,
            "failed_pages": len(images) - successful_pages,
            "total_characters": len(total_text),
            "total_tokens_used": total_tokens_used,
            "extracted_text": complete_text,
            "output_files": {
                "word_document": str(word_file),
                "complete_text": str(complete_file),
                "individual_pages": [str(output_path / f"page_{i:03d}.txt") for i in range(1, len(images) + 1)]
            },
            "page_results": extraction_results,
            "processing_time_ms": int(total_time),
            "status": final_status
        }

        # Save results summary
        summary_file = output_path / f"{Path(pdf_path).stem}_summary.json"
        with open(summary_file, 'w', encoding='utf-8') as f:
            json.dump(summary, f, indent=2, default=str)

        print(f"\n{'='*60}")
        print(f"EXTRACTION COMPLETE!")
        print(f"{'='*60}")
        print(f"✅ Successfully processed: {successful_pages}/{len(images)} pages")
        print(f"📝 Total characters extracted: {len(total_text):,}")
        print(f"🎯 Total tokens used: {total_tokens_used:,}")
        print(f"📁 Output directory: {output_path}")
        print(f"📄 Word document: {word_file}")
        print(f"📄 Database document ID: {document_id}")

        return summary


# Note: FastAPI app and endpoints have been moved to main.py to avoid conflicts
# This file now only contains the DatabaseOCR class for import

# Usage example for testing
def main():
    """Main function for testing"""
    # Initialize OCR processor
    ocr = DatabaseOCR()

    # Test API connection
    api_test = ocr.test_api_connection()
    if not api_test["success"]:
        print(f"❌ Mistral API connection failed: {api_test['error']}")
        return
    print("✅ Mistral API connection successful!")

    # Test database connection
    db_test = ocr.test_database_connection()
    if not db_test["success"]:
        print(f"❌ Database connection failed: {db_test['error']}")
        return
    print("✅ Database connection successful!")

    print("🚀 OCR Database API with Mistral Pixtral is ready!")
    print("📡 Start the FastAPI server with: uvicorn main:app --host 0.0.0.0 --port 8000 --reload")
    print("🌐 Frontend should connect to: http://localhost:8000")

if __name__ == "__main__":
    main()